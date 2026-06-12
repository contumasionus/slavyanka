#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерация большого Приложения А (Листинг) — 30+ страниц кода"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = 'Диплом_Готовый'

def setup_page(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

def add_frame(doc):
    sectPr = doc.sections[0]._sectPr
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '</w:pgBorders>'
    )
    sectPr.append(pgBorders)

def set_normal(doc):
    s = doc.styles['Normal']
    f = s.font; f.name = 'Times New Roman'; f.size = Pt(14)
    s.paragraph_format.line_spacing = 1.5

def heading(doc, text, size=16):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = True

def subh(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

def caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def code_block(doc, code):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(code); r.font.name = 'Courier New'; r.font.size = Pt(9)

def empty(doc):
    p = doc.add_paragraph(); r = p.add_run(''); r.font.size = Pt(10)

def page_br(doc):
    doc.add_page_break()

# ============================================================
doc = Document()
setup_page(doc); set_normal(doc); add_frame(doc)

heading(doc, 'Приложение А', 16)
heading(doc, 'Листинг программы', 14)
empty(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run('В данном приложении представлены ключевые фрагменты исходного кода веб-сервиса «Славянка». Листинги содержат основные модули клиентской и серверной частей приложения: сервер на Fastify, схему базы данных Prisma, хранилища состояния Pinia, API-клиенты, маршруты, контроллеры, сервисы и компоненты пользовательского интерфейса.')
r.font.name = 'Times New Roman'; r.font.size = Pt(14)

LISTINGS = [
("А.1", "Точка входа сервера (server.ts)", "Главный файл сервера Fastify", """import dotenv from 'dotenv';
import path from 'path';
dotenv.config({ path: path.resolve(__dirname, '../.env') });

import Fastify from 'fastify';
import cors from '@fastify/cors';
import jwt from '@fastify/jwt';
import swagger from '@fastify/swagger';
import swaggerUi from '@fastify/swagger-ui';
import multipart from '@fastify/multipart';
import helmet from '@fastify/helmet';
import rateLimit from '@fastify/rate-limit';
import { registerRoutes } from './routes';

const PORT = parseInt(process.env.PORT || '3001', 10);
const HOST = process.env.HOST || '0.0.0.0';

async function main() {
  // Создание экземпляра сервера с логированием
  const server = Fastify({ logger: { level: 'info' } });
  
  // Плагин CORS для кросс-доменных запросов
  await server.register(cors, { origin: true, credentials: true });
  
  // JWT для аутентификации
  await server.register(jwt, { 
    secret: process.env.JWT_SECRET || 'slavyanka-super-secret-key-2026' 
  });
  
  // Swagger для документации API
  await server.register(swagger, {
    openapi: {
      info: { 
        title: 'Славянка API', 
        description: 'API для магазина "Славянка"', 
        version: '1.0.0' 
      },
    },
  });
  await server.register(swaggerUi, { routePrefix: '/docs' });
  
  // Mulipart для загрузки файлов (до 10MB)
  await server.register(multipart, { 
    limits: { fileSize: 10 * 1024 * 1024 } 
  });
  
  // Helmet для защиты HTTP-заголовков
  await server.register(helmet, { contentSecurityPolicy: false });
  
  // Rate limit — не более 100 запросов в минуту
  await server.register(rateLimit, { 
    max: 100, 
    timeWindow: '1 minute' 
  });
  
  // Регистрация всех маршрутов приложения
  await registerRoutes(server);
  
  // Health check endpoint
  server.get('/health', async () => ({ status: 'ok' }));
  
  // Запуск сервера
  try {
    await server.listen({ port: PORT, host: HOST });
    console.log(`Сервер запущен на http://${HOST}:${PORT}`);
    console.log(`Swagger документация: http://localhost:${PORT}/docs`);
  } catch (err) {
    server.log.error(err);
    process.exit(1);
  }
}

main();"""),

("А.2", "Регистрация маршрутов API (routes/index.ts)", "Подключение всех маршрутов сервера", """import { FastifyInstance } from 'fastify';
import { authRoutes } from './auth.routes';
import { productsRoutes } from './products.routes';
import { categoriesRoutes } from './categories.routes';
import { ordersRoutes } from './orders.routes';
import { usersRoutes } from './users.routes';
import { notificationsRoutes } from './notifications.routes';
import { visitsRoutes } from './visits.routes';
import { uploadRoutes } from './upload.routes';
import { chatRoutes } from './chat.routes';
import { reviewsRoutes } from './reviews.routes';
import { newsRoutes } from './news.routes';
import { promoRoutes } from './promo.routes';
import { contactMessagesRoutes } from './contact-messages.routes';

/**
 * Регистрирует все маршруты API с соответствующими префиксами
 */
export async function registerRoutes(server: FastifyInstance) {
  server.register(authRoutes, { prefix: '/api/auth' });
  server.register(productsRoutes, { prefix: '/api/products' });
  server.register(categoriesRoutes, { prefix: '/api/categories' });
  server.register(ordersRoutes, { prefix: '/api/orders' });
  server.register(usersRoutes, { prefix: '/api/users' });
  server.register(notificationsRoutes, { prefix: '/api/notifications' });
  server.register(visitsRoutes, { prefix: '/api/visits' });
  server.register(uploadRoutes, { prefix: '/api' });
  server.register(chatRoutes, { prefix: '/api' });
  server.register(reviewsRoutes, { prefix: '/api/reviews' });
  server.register(newsRoutes, { prefix: '/api/news' });
  server.register(promoRoutes, { prefix: '/api/promo' });
  server.register(contactMessagesRoutes, { prefix: '/api' });
}"""),

("А.3", "Схема базы данных (schema.prisma)", "Модели данных Prisma", """// Генератор клиента для работы с БД
generator client { provider = "prisma-client-js" }

// Источник данных: SQLite (файл dev.db)
datasource db { provider = "sqlite" url = "file:./dev.db" }

// Модель пользователя
model User {
  id            String   @id @default(uuid())
  email         String   @unique        // Уникальный email
  password      String                  // bcrypt hash
  name          String                  // Имя пользователя
  role          String   @default("customer") // customer | admin
  verified      Boolean  @default(false)      // Подтверждён ли email
  createdAt     DateTime @default(now())
  updatedAt     DateTime @updatedAt
  orders        Order[]                  // Связь с заказами
  reviews       Review[]                 // Связь с отзывами
  chatSessions  ChatSession[]            // Связь с чатами
  viewedProducts ViewedProduct[]
  @@map("users")
}

// Код подтверждения email
model VerificationCode {
  id        String   @id @default(uuid())
  email     String
  code      String                       // 6-значный код
  expiresAt DateTime                     // Срок действия
  createdAt DateTime @default(now())
  @@index([email, code])
  @@map("verification_codes")
}

// Категория товаров
model Category {
  id          String    @id @default(uuid())
  name        String                     // Название категории
  slug        String    @unique           // URL-адрес категории
  description String?
  imageUrl    String?
  createdAt   DateTime  @default(now())
  updatedAt   DateTime  @updatedAt
  products    Product[]
  @@map("categories")
}

// Товар
model Product {
  id            String    @id @default(uuid())
  name          String                   // Название товара
  description   String?
  price         Float                    // Цена в рублях
  weight        String?                  // Вес (например, '500 г')
  imageUrl      String?
  categoryId    String                   // Внешний ключ к категории
  inStock       Boolean   @default(true) // Наличие
  composition   String?                  // Состав
  manufacturer  String?                  // Производитель
  country       String?                  // Страна производства
  shelfLife     String?                  // Срок годности
  calories      Int?                     // Калорийность
  proteins      Float?                   // Белки
  fats          Float?                   // Жиры
  carbohydrates Float?                   // Углеводы
  avgRating     Float     @default(0)    // Средний рейтинг
  reviewCount   Int       @default(0)    // Количество отзывов
  discountPrice Float?                   // Цена со скидкой
  discountUntil DateTime?
  isPromo       Boolean   @default(false) // Акционный товар
  deliveryDate  DateTime?
  createdAt     DateTime  @default(now())
  updatedAt     DateTime  @updatedAt
  category      Category  @relation(fields: [categoryId], references: [id])
  reviews       Review[]
  orderItems    OrderItem[]
  @@map("products")
}

// Заказ
model Order {
  id              String     @id @default(uuid())
  userId          String
  status          String     @default("new") // new | processing | completed | cancelled
  totalAmount     Float
  customerName    String     @default("")
  customerPhone   String     @default("")
  customerAddress String     @default("")
  comment         String     @default("")
  deliveryType    String     @default("pickup") // pickup | delivery
  deliveryTime    String?
  deliveryAddress String?
  paymentMethod   String     @default("")
  cardLast4       String     @default("")
  paid            Boolean    @default(false)
  createdAt       DateTime   @default(now())
  updatedAt       DateTime   @updatedAt
  user            User       @relation(fields: [userId], references: [id])
  items           OrderItem[]
  @@map("orders")
}

// Позиция заказа
model OrderItem {
  id        String  @id @default(uuid())
  orderId   String
  productId String
  quantity  Int
  price     Float
  product   Product @relation(fields: [productId], references: [id])
  order     Order   @relation(fields: [orderId], references: [id])
  @@map("order_items")
}

// Отзыв
model Review {
  id        String   @id @default(uuid())
  userId    String
  productId String
  rating    Int                          // Оценка 1-5
  text      String?
  photo     String?
  createdAt DateTime @default(now())
  updatedAt DateTime @updatedAt
  user      User     @relation(fields: [userId], references: [id])
  product   Product  @relation(fields: [productId], references: [id])
  @@unique([userId, productId])          // Один отзыв на товар от пользователя
  @@map("reviews")
}

// Новость
model News {
  id        String   @id @default(uuid())
  title     String
  slug      String   @unique
  content   String
  imageUrl  String?
  excerpt   String?
  tags      String?
  published Boolean  @default(false)
  createdAt DateTime @default(now())
  updatedAt DateTime @updatedAt
  @@map("news")
}

// Просмотренный товар
model ViewedProduct {
  id        String   @id @default(uuid())
  userId    String
  productId String
  viewedAt  DateTime @default(now())
  user      User     @relation(fields: [userId], references: [id])
  product   Product  @relation(fields: [productId], references: [id])
  @@index([userId, viewedAt])
  @@map("viewed_products")
}

// Сессия чата
model ChatSession {
  id        String        @id @default(uuid())
  userId    String?
  guestId   String        // ID гостя (если без авторизации)
  userName  String?
  status    String        @default("active") // active | closed
  createdAt DateTime      @default(now())
  updatedAt DateTime      @updatedAt
  messages  ChatMessage[]
  user      User?         @relation(fields: [userId], references: [id])
  @@map("chat_sessions")
}

// Сообщение чата
model ChatMessage {
  id         String       @id @default(uuid())
  sessionId  String
  senderType String       // 'user' или 'admin'
  content    String
  attachment String?
  isRead     Boolean      @default(false)
  readAt     DateTime?
  createdAt  DateTime     @default(now())
  session    ChatSession  @relation(fields: [sessionId], references: [id], onDelete: Cascade)
  @@map("chat_messages")
}

// Промокод
model PromoCode {
  id        String   @id @default(uuid())
  email     String
  code      String   @unique
  used      Boolean  @default(false)
  createdAt DateTime @default(now())
  expiresAt DateTime
  @@map("promo_codes")
}

// Сообщение обратной связи
model ContactMessage {
  id         String    @id @default(uuid())
  name       String
  email      String
  phone      String    @default("")
  topic      String    @default("")
  message    String
  adminReply String?
  repliedAt  DateTime?
  isRead     Boolean   @default(false)
  createdAt  DateTime  @default(now())
  updatedAt  DateTime  @updatedAt
  @@map("contact_messages")
}

// Журнал аудита
model AuditLog {
  id        String   @id @default(uuid())
  userId    String?
  action    String                   // Действие: login, create_order, update_product
  entity    String?                  // Сущность: order, product, user
  entityId  String?                  // ID сущности
  details   String?                  // Дополнительные детали
  ipAddress String?
  userAgent String?
  createdAt DateTime @default(now())
  user      User?    @relation(fields: [userId], references: [id], onDelete: SetNull)
  @@map("audit_logs")
}"""),

("А.4", "Хранилище корзины (cart.store.ts)", "Pinia store для управления корзиной", """import { defineStore } from 'pinia';
import { ref, computed } from 'vue';

// Хранилище корзины с использованием Pinia Composition API
export const useCartStore = defineStore('cart', () => {
  const items = ref([]);

  // Вычисляемые свойства
  const totalItems = computed(() => 
    items.value.reduce((sum, item) => sum + item.quantity, 0)
  );
  const totalPrice = computed(() => 
    items.value.reduce((sum, item) => sum + item.price * item.quantity, 0)
  );

  // Добавление товара в корзину
  function addItem(product) {
    const existing = items.value.find(i => i.id === product.id);
    if (existing) {
      existing.quantity++;
    } else {
      items.value.push({ ...product, quantity: 1 });
    }
  }

  // Удаление товара из корзины
  function removeItem(productId) {
    items.value = items.value.filter(i => i.id !== productId);
  }

  // Изменение количества товара
  function updateQuantity(productId, qty) {
    const item = items.value.find(i => i.id === productId);
    if (item) {
      item.quantity = Math.max(1, qty);
    }
  }

  // Очистка корзины
  function clearCart() {
    items.value = [];
  }

  return { items, totalItems, totalPrice, addItem, removeItem, updateQuantity, clearCart };
});

// Функции для синхронизации корзины с auth store
let cartStorageKey = 'slavyanka-cart-guest';
let cartUserId = null;

export function setCartUserId(userId) {
  cartUserId = userId;
  cartStorageKey = userId ? `slavyanka-cart-${userId}` : 'slavyanka-cart-guest';
}

export function resetCart() {
  // Сброс корзины в памяти
}"""),

("А.5", "Хранилище авторизации (auth.store.ts)", "Pinia store для аутентификации", """import { defineStore } from 'pinia';
import { ref, computed } from 'vue';
import { authApi } from '../api/auth.api';

// Вспомогательные функции для работы с localStorage
function getLocalStorageItem(key) {
  try { return localStorage.getItem(key); }
  catch { return null; }
}
function setLocalStorageItem(key, value) {
  try { localStorage.setItem(key, value); }
  catch (error) { console.error('localStorage недоступен:', error); }
}
function removeLocalStorageItem(key) {
  try { localStorage.removeItem(key); }
  catch (error) { console.error('localStorage недоступен:', error); }
}

export const useAuthStore = defineStore('auth', () => {
  const token = ref(getLocalStorageItem('token'));
  const user = ref(null);

  // Вычисляемые свойства
  const isAuthenticated = computed(() => !!token.value);
  const isAdmin = computed(() => user.value?.role === 'admin');

  // Вход в систему
  async function login(email, password) {
    const response = await authApi.login(email, password);
    token.value = response.token;
    user.value = response.user;
    setLocalStorageItem('token', response.token);
    // Переключаем корзину и избранное на этого пользователя
    setCartUserId(response.user.id);
    setFavoritesUserId(response.user.id);
  }

  // Регистрация нового пользователя
  async function register(email, password, name) {
    const response = await authApi.register(email, password, name);
    token.value = response.token;
    user.value = response.user;
    setLocalStorageItem('token', response.token);
    setCartUserId(response.user.id);
    setFavoritesUserId(response.user.id);
  }

  // Загрузка данных пользователя
  async function fetchUser() {
    if (!token.value) return;
    user.value = await authApi.getMe();
    setCartUserId(user.value.id);
    setFavoritesUserId(user.value.id);
  }

  // Выход из системы
  function logout() {
    token.value = null;
    user.value = null;
    removeLocalStorageItem('token');
    // Очистка данных чата при выходе
    ['chat-session-id','chat-guest-id','chat-messages','chat-user-hash']
      .forEach(key => removeLocalStorageItem(key));
    resetCart();
    resetFavorites();
    setCartUserId(null);
    setFavoritesUserId(null);
  }

  return { token, user, isAuthenticated, isAdmin, login, register, fetchUser, logout };
});"""),

("А.6", "Хранилище избранного (favorites.store.ts)", "Модуль избранного с localStorage", """import { ref, computed, watch } from 'vue';

const GUEST_STORAGE_KEY = 'slavyanka-favorites-guest';
const STORAGE_VERSION_KEY = 'slavyanka-favorites-version';
const STORAGE_VERSION = 2;
let storageKey = GUEST_STORAGE_KEY;

// Интерфейс элемента избранного
export interface FavoriteItem {
  id: string;
  name: string;
  price: number;
  imageUrl?: string;
  inStock: boolean;
  description?: string;
  composition?: string;
  manufacturer?: string;
  country?: string;
  shelfLife?: string;
  weight?: string;
  calories?: number;
  proteins?: number;
  fats?: number;
  carbohydrates?: number;
  discountPrice?: number;
  discountUntil?: string;
  isPromo?: boolean;
  avgRating?: number;
  reviewCount?: number;
}

const favorites = ref([]);

// Загрузка избранного из localStorage
function loadFavorites() {
  try {
    const saved = localStorage.getItem(storageKey);
    favorites.value = saved ? JSON.parse(saved) : [];
  } catch (e) {
    favorites.value = [];
  }
}

// Сохранение избранного в localStorage
function saveFavorites() {
  try {
    localStorage.setItem(storageKey, JSON.stringify(favorites.value));
  } catch (e) {
    console.error('Failed to save favorites:', e);
  }
}

loadFavorites();

// Автосохранение при изменениях
watch(favorites, () => { saveFavorites(); }, { deep: true });

// Сброс избранного
export function resetFavorites() { favorites.value = []; }

// Привязка к userId
export function setFavoritesUserId(userId) {
  const newKey = userId ? `slavyanka-favorites-${userId}` : GUEST_STORAGE_KEY;
  if (newKey !== storageKey) {
    storageKey = newKey;
    loadFavorites();
  }
}

// Хук для использования в компонентах
export function useFavorites() {
  const count = computed(() => favorites.value.length);

  function toggle(item) {
    const index = favorites.value.findIndex(f => f.id === item.id);
    if (index === -1) favorites.value.push(item);
    else favorites.value.splice(index, 1);
  }

  function isFavorite(id) { return favorites.value.some(f => f.id === id); }
  function remove(id) {
    const index = favorites.value.findIndex(f => f.id === id);
    if (index !== -1) favorites.value.splice(index, 1);
  }
  function getAll() { return [...favorites.value]; }

  return { favorites, count, toggle, isFavorite, remove, getAll };
}"""),

("А.7", "API-клиент для авторизации (auth.api.ts)", "Клиент для взаимодействия с сервером авторизации", """import apiClient from './client';

// API для работы с аутентификацией
export const authApi = {
  // Вход в систему
  async login(email, password) {
    const response = await apiClient.post('/auth/login', { email, password });
    return response.data;
  },

  // Регистрация нового пользователя
  async register(email, password, name) {
    const response = await apiClient.post('/auth/register', { email, password, name });
    return response.data;
  },

  // Получение информации о текущем пользователе
  async getMe() {
    const response = await apiClient.get('/auth/me');
    return response.data;
  },

  // Отправка кода подтверждения на email
  async sendVerification(email) {
    const response = await apiClient.post('/auth/send-verification', { email });
    return response.data;
  },

  // Проверка кода подтверждения
  async verifyCode(email, code) {
    const response = await apiClient.post('/auth/verify-code', { email, code });
    return response.data;
  },
};"""),

("А.8", "API-клиент для товаров (products.api.ts)", "Клиент для работы с каталогом товаров", """import apiClient from './client';

export const productsApi = {
  // Получение списка товаров с фильтрацией
  async getAll(params = {}) {
    const response = await apiClient.get('/products', { params });
    return response.data;
  },

  // Получение товара по ID
  async getById(id) {
    const response = await apiClient.get(`/products/${id}`);
    return response.data;
  },

  // Создание товара (админ)
  async create(data) {
    const response = await apiClient.post('/products', data);
    return response.data;
  },

  // Обновление товара (админ)
  async update(id, data) {
    const response = await apiClient.put(`/products/${id}`, data);
    return response.data;
  },

  // Удаление товара (админ)
  async delete(id) {
    const response = await apiClient.delete(`/products/${id}`);
    return response.data;
  },
};"""),

("А.9", "Валидация форм (validation.ts)", "Функции валидации и защита от XSS", """// Валидация email-адреса
export function validateEmail(email) {
  return /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/.test(email);
}

// Валидация номера телефона (Россия/СНГ)
export function validatePhone(phone) {
  return /^[+]?[\\d\\s()-]{10,15}$/.test(phone);
}

// Проверка, что поле не пустое
export function validateRequired(value) {
  return value.trim().length > 0;
}

// Проверка минимальной длины
export function validateMinLength(value, min) {
  return value.trim().length >= min;
}

// Валидация цены (положительное число)
export function validatePrice(price) {
  const num = parseFloat(price);
  return !isNaN(num) && num > 0;
}

// Санитизация ввода для защиты от XSS-атак
export function sanitizeInput(input) {
  return input
    .replace(/&/g, '&')
    .replace(/</g, '<')
    .replace(/>/g, '>')
    .replace(/"/g, '"')
    .replace(/'/g, '&#x27;')
    .replace(/\\//g, '&#x2F;');
}"""),

("А.10", "Хеширование паролей (hash.ts)", "Утилита для хеширования паролей (bcrypt)", """import bcrypt from 'bcryptjs';

const SALT_ROUNDS = 10;

// Хеширование пароля перед сохранением в БД
export async function hashPassword(password) {
  return bcrypt.hash(password, SALT_ROUNDS);
}

// Сравнение введённого пароля с хешем из БД
export async function comparePassword(password, hash) {
  return bcrypt.compare(password, hash);
}"""),

("А.11", "Маршруты авторизации (auth.routes.ts)", "Обработка HTTP-запросов авторизации", """import { FastifyInstance } from 'fastify';
import { authController } from '../controllers/auth.controller';

export async function authRoutes(server) {
  // Регистрация нового пользователя
  // POST /api/auth/register { email, password, name }
  server.post('/register', authController.register);
  
  // Вход в систему
  // POST /api/auth/login { email, password }
  server.post('/login', authController.login);
  
  // Получение текущего пользователя (требуется JWT-токен)
  // GET /api/auth/me
  server.get('/me', { preHandler: [server.authenticate] }, authController.getMe);
  
  // Отправка кода подтверждения на email
  // POST /api/auth/send-verification { email }
  server.post('/send-verification', authController.sendVerificationCode);
  
  // Проверка кода подтверждения
  // POST /api/auth/verify-code { email, code }
  server.post('/verify-code', authController.verifyCode);
}"""),

("А.12", "Сервис авторизации (auth.service.ts)", "Бизнес-логика аутентификации", """import { PrismaClient } from '@prisma/client';
import { hashPassword, comparePassword } from '../utils/hash';

const prisma = new PrismaClient();

export const authService = {
  // Регистрация нового пользователя
  async register({ email, password, name }) {
    // Проверка, не занят ли email
    const existing = await prisma.user.findUnique({ where: { email } });
    if (existing) throw new Error('Email уже используется');
    
    // Хеширование пароля
    const hashedPassword = await hashPassword(password);
    
    // Создание пользователя в БД
    const user = await prisma.user.create({
      data: { email, password: hashedPassword, name },
    });
    return { id: user.id, email: user.email, name: user.name, role: user.role };
  },

  // Вход в систему
  async login({ email, password }) {
    // Поиск пользователя по email
    const user = await prisma.user.findUnique({ where: { email } });
    if (!user) throw new Error('Неверный email или пароль');
    
    // Сравнение паролей
    const valid = await comparePassword(password, user.password);
    if (!valid) throw new Error('Неверный email или пароль');
    
    return { id: user.id, email: user.email, name: user.name, role: user.role };
  },
  
  // Получение пользователя по ID
  async getUserById(id) {
    const user = await prisma.user.findUnique({
      where: { id },
      select: { id: true, email: true, name: true, role: true, verified: true, createdAt: true },
    });
    if (!user) throw new Error('Пользователь не найден');
    return user;
  },
};"""),

("А.13", "Контроллер авторизации (auth.controller.ts)", "Обработчик запросов авторизации", """import { FastifyRequest, FastifyReply } from 'fastify';
import { AuthService } from '../services/auth.service';
import { VerificationService } from '../services/verification.service';

export class AuthController {
  constructor(
    private authService: AuthService,
    private verificationService: VerificationService,
  ) {}

  // Вход в систему
  async login(request, reply) {
    try {
      const user = await this.authService.login(request.body);
      // Генерация JWT-токена
      const token = request.server.jwt.sign({ 
        id: user.id, email: user.email, role: user.role 
      });
      reply.send({ token, user });
    } catch (error) {
      reply.status(401).send({ error: error.message });
    }
  }

  // Регистрация
  async register(request, reply) {
    try {
      const user = await this.authService.register(request.body);
      const token = request.server.jwt.sign({ 
        id: user.id, email: user.email, role: user.role 
      });
      reply.status(201).send({ token, user });
    } catch (error) {
      reply.status(400).send({ error: error.message });
    }
  }

  // Получение профиля
  async getMe(request, reply) {
    try {
      const user = request.user;
      const userData = await this.authService.getUserById(user.id);
      reply.send(userData);
    } catch (error) {
      reply.status(404).send({ error: error.message });
    }
  }

  // Отправка кода подтверждения
  async sendVerificationCode(request, reply) {
    try {
      const { email } = request.body;
      // Проверка, не занят ли email
      const existingUser = await prisma.user.findUnique({ where: { email } });
      if (existingUser) {
        reply.status(400).send({ error: 'Email уже зарегистрирован' });
        return;
      }
      await this.verificationService.sendCode(email);
      reply.send({ success: true, message: 'Код отправлен на email' });
    } catch (error) {
      reply.status(500).send({ error: 'Ошибка отправки кода' });
    }
  }

  // Проверка кода подтверждения
  async verifyCode(request, reply) {
    try {
      const { email, code } = request.body;
      const isValid = await this.verificationService.verifyCode(email, code);
      if (!isValid) {
        reply.status(400).send({ error: 'Неверный или истёкший код' });
        return;
      }
      reply.send({ success: true });
    } catch (error) {
      reply.status(500).send({ error: 'Ошибка проверки кода' });
    }
  }
}"""),

("А.14", "Сервис заказов (orders.service.ts)", "Бизнес-логика обработки заказов", """import { PrismaClient } from '@prisma/client';

const prisma = new PrismaClient();

export const ordersService = {
  // Получение списка заказов с пагинацией
  async getAll({ userId, status, page = 1, limit = 20 }) {
    const where = {};
    if (userId) where.userId = userId;
    if (status) where.status = status;

    const [items, total] = await Promise.all([
      prisma.order.findMany({
        where,
        include: { items: { include: { product: true } } },
        skip: (page - 1) * limit,
        take: limit,
        orderBy: { createdAt: 'desc' },
      }),
      prisma.order.count({ where }),
    ]);
    return { items, total, page, limit };
  },

  // Получение заказа по ID
  async getById(id) {
    return prisma.order.findUnique({
      where: { id },
      include: { items: { include: { product: true } } },
    });
  },

  // Создание нового заказа
  async create(data) {
    const { 
      userId, items, customerName, customerPhone, 
      deliveryAddress, deliveryType, paymentMethod, comment 
    } = data;
    
    // Расчёт общей суммы
    const totalAmount = items.reduce(
      (sum, item) => sum + item.price * item.quantity, 0
    );

    return prisma.order.create({
      data: {
        userId,
        totalAmount,
        customerName,
        customerPhone,
        deliveryAddress,
        deliveryType,
        paymentMethod,
        comment,
        items: {
          create: items.map(item => ({
            productId: item.productId,
            quantity: item.quantity,
            price: item.price,
          })),
        },
      },
      include: { items: true },
    });
  },

  // Обновление статуса заказа
  async updateStatus(id, status) {
    return prisma.order.update({ where: { id }, data: { status } });
  },
};"""),

("А.15", "Контроллер товаров (products.controller.ts)", "Обработчик API-запросов к каталогу", """import { productsService } from '../services/products.service';

export const productsController = {
  // GET /api/products — список товаров
  async getAll(request, reply) {
    const { category, search, sort, page, limit } = request.query;
    const products = await productsService.getAll({
      category, search, sort, page: parseInt(page) || 1, 
      limit: parseInt(limit) || 20,
    });
    return products;
  },

  // GET /api/products/:id — один товар
  async getById(request, reply) {
    const { id } = request.params;
    const product = await productsService.getById(id);
    if (!product) {
      return reply.status(404).send({ error: 'Товар не найден' });
    }
    return product;
  },

  // POST /api/products — создание (админ)
  async create(request, reply) {
    const product = await productsService.create(request.body);
    return reply.status(201).send(product);
  },

  // PUT /api/products/:id — обновление (админ)
  async update(request, reply) {
    const { id } = request.params;
    const product = await productsService.update(id, request.body);
    return product;
  },

  // DELETE /api/products/:id — удаление (админ)
  async delete(request, reply) {
    const { id } = request.params;
    await productsService.delete(id);
    return reply.status(204).send();
  },
};"""),

("А.16", "Сервис товаров (products.service.ts)", "Бизнес-логика работы с товарами", """import { PrismaClient } from '@prisma/client';

const prisma = new PrismaClient();

export const productsService = {
  // Получение списка с фильтрацией и пагинацией
  async getAll({ category, search, sort, page, limit }) {
    const where = {};
    
    // Фильтр по категории
    if (category) {
      where.category = { slug: category };
    }
    
    // Поиск по названию
    if (search) {
      where.name = { contains: search };
    }

    // Сортировка
    let orderBy = { createdAt: 'desc' };
    if (sort === 'price_asc') orderBy = { price: 'asc' };
    if (sort === 'price_desc') orderBy = { price: 'desc' };
    if (sort === 'name') orderBy = { name: 'asc' };
    if (sort === 'rating') orderBy = { avgRating: 'desc' };

    const [items, total] = await Promise.all([
      prisma.product.findMany({
        where,
        include: { category: true },
        skip: (page - 1) * limit,
        take: limit,
        orderBy,
      }),
      prisma.product.count({ where }),
    ]);
    
    return { items, total, page, limit };
  },

  async getById(id) {
    return prisma.product.findUnique({
      where: { id },
      include: { 
        category: true,
        reviews: { include: { user: true }, orderBy: { createdAt: 'desc' } },
      },
    });
  },

  async create(data) {
    return prisma.product.create({ data });
  },

  async update(id, data) {
    return prisma.product.update({ where: { id }, data });
  },

  async delete(id) {
    return prisma.product.delete({ where: { id } });
  },
};"""),

("А.17", "Контроллер чата (chat.controller.ts)", "Обработка WebSocket-чата поддержки", """import { PrismaClient } from '@prisma/client';

const prisma = new PrismaClient();

export const chatController = {
  // Создание новой сессии чата
  async createSession(request, reply) {
    const { userId, guestId, userName } = request.body;
    const session = await prisma.chatSession.create({
      data: { userId, guestId, userName },
    });
    return session;
  },

  // Получение сообщений сессии
  async getMessages(request, reply) {
    const { sessionId } = request.params;
    const messages = await prisma.chatMessage.findMany({
      where: { sessionId },
      orderBy: { createdAt: 'asc' },
    });
    return messages;
  },

  // Отправка сообщения
  async sendMessage(request, reply) {
    const { sessionId, senderType, content } = request.body;
    const message = await prisma.chatMessage.create({
      data: { sessionId, senderType, content },
    });
    return message;
  },

  // Получение активных сессий (для админа)
  async getActiveSessions(request, reply) {
    const sessions = await prisma.chatSession.findMany({
      where: { status: 'active' },
      include: { 
        messages: { take: 1, orderBy: { createdAt: 'desc' } },
        user: { select: { id: true, name: true } },
      },
      orderBy: { updatedAt: 'desc' },
    });
    return sessions;
  },
};"""),

("А.18", "Главный компонент приложения (App.vue)", "Корневой Vue-компонент", """<template>
  <div id="app">
    <!-- Пользовательская часть -->
    <template v-if="!isAdminRoute">
      <AppHeader />
      <main class="main">
        <router-view v-slot="{ Component, route }">
          <transition name="page" mode="out-in">
            <component :is="Component" :key="route.path" />
          </transition>
        </router-view>
      </main>
      <AppFooter />
      <ScrollToTop />
      <ProductModal />
      <ChatWidget />
    </template>
    <!-- Административная панель -->
    <template v-else>
      <router-view />
    </template>
  </div>
</template>

<script setup lang="ts">
import { computed, onMounted } from 'vue';
import { useRoute } from 'vue-router';
import { useAuthStore } from './stores/auth.store';
import AppHeader from './components/layout/AppHeader.vue';
import AppFooter from './components/layout/AppFooter.vue';
import ScrollToTop from './components/ui/ScrollToTop.vue';
import ChatWidget from './components/chat/ChatWidget.vue';
import ProductModal from './components/product/ProductModal.vue';

const route = useRoute();
const authStore = useAuthStore();
const isAdminRoute = computed(() => route.path.startsWith('/admin'));

onMounted(async () => {
  // Загрузка пользователя при наличии токена
  if (authStore.token) {
    await authStore.fetchUser();
  }
});
</script>

<style>
#app { min-height: 100vh; display: flex; flex-direction: column; }
.main { flex: 1; padding-top: 80px; }
@media (max-width: 768px) { .main { padding-top: 64px; } }
@media (max-width: 480px) { .main { padding-top: 60px; } }
</style>"""),

("А.19", "Маршрутизация (router/index.ts)", "Конфигурация Vue Router", """import { createRouter, createWebHistory } from 'vue-router';
import { useAuthStore } from '../stores/auth.store';

const routes = [
  { path: '/', component: () => import('../views/HomePage.vue') },
  { path: '/catalog', component: () => import('../views/CatalogPage.vue') },
  { path: '/promo', component: () => import('../views/PromoPage.vue') },
  { path: '/news', component: () => import('../views/NewsPage.vue') },
  { path: '/news/:slug', component: () => import('../views/NewsArticlePage.vue') },
  { path: '/about', component: () => import('../views/AboutPage.vue') },
  { path: '/contacts', component: () => import('../views/ContactsPage.vue') },
  { path: '/cart', component: () => import('../views/CartPage.vue') },
  { path: '/privacy', component: () => import('../views/PrivacyPage.vue') },
  { path: '/login', component: () => import('../views/LoginPage.vue') },
  { path: '/register', component: () => import('../views/RegisterPage.vue') },
  { path: '/profile', component: () => import('../views/ProfilePage.vue'), 
    meta: { requiresAuth: true } },
  { path: '/admin/login', component: () => import('../views/admin/LoginPage.vue') },
  { 
    path: '/admin', component: () => import('../views/admin/AdminLayout.vue'),
    children: [
      { path: '', redirect: '/admin/dashboard' },
      { path: 'dashboard', component: () => import('../views/admin/DashboardPage.vue') },
      { path: 'products', component: () => import('../views/admin/ProductsPage.vue') },
      { path: 'categories', component: () => import('../views/admin/CategoriesPage.vue') },
      { path: 'orders', component: () => import('../views/admin/OrdersPage.vue') },
      { path: 'users', component: () => import('../views/admin/UsersPage.vue') },
      { path: 'chats', component: () => import('../views/admin/ChatsPage.vue') },
      { path: 'news', component: () => import('../views/admin/AdminNewsPage.vue') },
      { path: 'messages', component: () => import('../views/admin/ContactMessagesPage.vue') },
    ],
  },
];

const router = createRouter({
  history: createWebHistory(),
  routes,
  scrollBehavior() { return { top: 0, behavior: 'smooth' }; },
});

// Guard для проверки авторизации
router.beforeEach(async (to, from, next) => {
  const authStore = useAuthStore();
  
  if (to.meta.requiresAuth && !authStore.isAuthenticated) {
    next('/login');
    return;
  }
  
  if (to.path.startsWith('/admin') && to.path !== '/admin/login') {
    if (!authStore.isAuthenticated) {
      next('/admin/login');
      return;
    }
    if (!authStore.isAdmin) {
      next('/');
      return;
    }
  }
  
  next();
});

export default router;"""),

("А.20", "Список населённых пунктов (nearbySettlements.ts)", "Данные для умного адреса", """// Интерфейс населённого пункта
export interface Settlement {
  name: string;    // Название
  type: string;    // Тип: village, selo, city, settlement
}

// Список ближайших населённых пунктов к магазину "Славянка" (д. Курилово)
export const nearbySettlements = [
  { name: 'д. Курилово', type: 'village' },
  { name: 'с. Ворша', type: 'selo' },
  { name: 'д. Устье', type: 'village' },
  { name: 'д. Березняки', type: 'village' },
  { name: 'с. Колокша', type: 'selo' },
  { name: 'д. Илевники', type: 'village' },
  { name: 'г. Собинка', type: 'city' },
  { name: 'д. Рождествено', type: 'village' },
  { name: 'д. Черкутино', type: 'village' },
  { name: 'д. Бабаево', type: 'village' },
  { name: 'п. Воровский', type: 'settlement' },
  { name: 'д. Епишиха', type: 'village' },
];

// Функция поиска населённых пунктов по строке запроса
export function searchSettlements(query) {
  if (!query || query.length < 2) return [];
  const lower = query.toLowerCase();
  return nearbySettlements.filter(s => 
    s.name.toLowerCase().includes(lower)
  );
}"""),

("А.21", "Сервис верификации email (verification.service.ts)", "Отправка и проверка кода подтверждения", """import { PrismaClient } from '@prisma/client';
import emailjs from '@emailjs/nodejs';

export class VerificationService {
  constructor(private prisma: PrismaClient) {}

  // Генерация 6-значного кода
  private generateCode() {
    return Math.floor(100000 + Math.random() * 900000).toString();
  }

  // Отправка кода на email через EmailJS
  async sendCode(email) {
    const code = this.generateCode();
    
    // Сохранение кода в БД со сроком 15 минут
    await this.prisma.verificationCode.create({
      data: {
        email,
        code,
        expiresAt: new Date(Date.now() + 15 * 60 * 1000),
      },
    });

    // Отправка через EmailJS
    await emailjs.send(
      process.env.EMAILJS_SERVICE_ID,
      process.env.EMAILJS_TEMPLATE_ID,
      { to_email: email, verification_code: code },
      { publicKey: process.env.EMAILJS_PUBLIC_KEY }
    );
  }

  // Проверка кода
  async verifyCode(email, code) {
    const record = await this.prisma.verificationCode.findFirst({
      where: { email, code, expiresAt: { gt: new Date() } },
    });
    return !!record;
  }
}"""),
]

for listing_id, title, desc, code in LISTINGS:
    empty(doc)
    subh(doc, f'{listing_id} {title}')
    caption(doc, f'Листинг {listing_id} – {desc}')
    empty(doc)
    code_block(doc, code)
    empty(doc)
    page_br(doc)

empty(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Полный текст программы веб-сервиса представлен в электронном виде и прилагается к дипломному проекту на CD-диске.')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

os.makedirs(os.path.join(OUTPUT_DIR, 'Приложения'), exist_ok=True)
doc.save(os.path.join(OUTPUT_DIR, 'Приложения', 'Приложение А (Листинг).docx'))
print(f'[OK] Приложение А (Листинг).docx — {sum(len(c) for _,_,_,c in LISTINGS)} символов кода')