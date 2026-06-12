#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генерация полного комплекта дипломных документов
для интернет-магазина продуктов «Славянка»
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, shutil, datetime

SOURCE_DIR = 'Диплом полностью'
OUTPUT_DIR = 'Диплом_Готовый'

def ensure_output_dir():
    """Создаёт структуру папок для вывода"""
    dirs = [
        OUTPUT_DIR,
        os.path.join(OUTPUT_DIR, 'Приложения'),
        os.path.join(OUTPUT_DIR, 'Карта сайта'),
    ]
    for d in dirs:
        os.makedirs(d, exist_ok=True)

def copy_static_files():
    """Копирует файлы, которые не нужно генерировать"""
    static_files = [
        ('Антиплагиат.pdf', 'Антиплагиат.pdf'),
        ('Презентация.pptx', 'Презентация.pptx'),
    ]
    for src_name, dst_name in static_files:
        src = os.path.join(SOURCE_DIR, src_name)
        dst = os.path.join(OUTPUT_DIR, dst_name)
        if os.path.exists(src) and not os.path.exists(dst):
            shutil.copy2(src, dst)
    
    # Карта сайта изображения
    site_map_src = os.path.join(SOURCE_DIR, 'Карта сайта')
    site_map_dst = os.path.join(OUTPUT_DIR, 'Карта сайта')
    if os.path.exists(site_map_src):
        for f in os.listdir(site_map_src):
            src_f = os.path.join(site_map_src, f)
            dst_f = os.path.join(site_map_dst, f)
            if os.path.isfile(src_f) and not os.path.exists(dst_f):
                shutil.copy2(src_f, dst_f)

# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ ФОРМАТИРОВАНИЯ
# ============================================================

def setup_page(doc, left_margin=3, other_margin=1.5):
    """Настройка полей страницы"""
    for section in doc.sections:
        section.top_margin = Cm(other_margin)
        section.bottom_margin = Cm(other_margin)
        section.left_margin = Cm(left_margin)
        section.right_margin = Cm(other_margin)

def add_frame(doc, section=None):
    """Добавляет рамку (обрамление страницы) как в ГОСТ"""
    if section is None:
        section = doc.sections[0]
    
    sectPr = section._sectPr
    # Создаём элемент pageBorders
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '</w:pgBorders>'
    )
    sectPr.append(pgBorders)

def set_normal_style(doc):
    """Устанавливает стандартный стиль"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

def add_heading_centered(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_paragraph(doc, text, bold=False, size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, first_line_indent=1.25):
    p = doc.add_paragraph()
    p.alignment = alignment
    if indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_empty_line(doc, size=10):
    p = doc.add_paragraph()
    run = p.add_run('')
    run.font.size = Pt(size)
    return p

def add_heading_section(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_subheading(doc, text, size=15, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_subsubheading(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def set_cell_text(cell, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_table_caption(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = True
    return p

def add_bullet_points(doc, items, indent_level=0):
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(indent_level * 1)
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def add_page_numbers(doc):
    """Добавляет номера страниц внизу"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Простой номер страницы
        run = p.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(12)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

# ============================================================
# ГЕНЕРАЦИЯ ТИТУЛЬНОГО ЛИСТА
# ============================================================
def generate_titulnik():
    doc = Document()
    setup_page(doc, left_margin=3, other_margin=1.5)
    set_normal_style(doc)
    
    add_empty_line(doc)
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Государственное автономное профессиональное образовательное учреждение\nВладимирской области\n«Владимирский политехнический колледж»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Специальность 09.02.07 Информационные системы и программирование')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    add_empty_line(doc)
    add_empty_line(doc)
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ДИПЛОМНЫЙ ПРОЕКТ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('РАЗРАБОТКА ВЕБ-СЕРВИСА ДЛЯ ИНТЕРНЕТ-МАГАЗИНА ПРОДУКТОВ «СЛАВЯНКА»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    add_empty_line(doc)
    add_empty_line(doc)
    add_empty_line(doc)
    
    # Подписи
    lines = [
        ('Студент', 'Г.М. Жуков'),
        ('Руководитель проекта', 'И.О. Меренцева'),
        ('Старший консультант', 'И.О. Меренцева'),
        ('Консультант по экономической части', 'В.В. Родионов'),
        ('Нормоконтроль', 'И.О. Меренцева'),
        ('Зав. отделением', 'В.Н. Степанова'),
    ]
    
    for role, name in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f'{role}\t\t\t\t\t{name}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    add_empty_line(doc)
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ПОЯСНИТЕЛЬНАЯ ЗАПИСКА\nДП 09.02.07 ИС-422.11.26 ПЗ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    add_empty_line(doc)
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    doc.save(os.path.join(OUTPUT_DIR, '1 Титульник.docx'))
    print('[OK] 1 Титульник.docx')

# ============================================================
# ГЕНЕРАЦИЯ ЗАДАНИЯ ДП
# ============================================================
def generate_zadanie():
    doc = Document()
    setup_page(doc, left_margin=2.5, other_margin=1.5)
    set_normal_style(doc)
    
    add_heading_centered(doc, 'ЗАДАНИЕ НА ДИПЛОМНОЕ ПРОЕКТИРОВАНИЕ', size=16)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Студент: Жуков Григорий Максимович', bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Специальность: 09.02.07 Информационные системы и программирование', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_paragraph(doc, 'Тема: Разработка веб-сервиса для интернет-магазина продуктов «Славянка»', bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Краткая характеристика: Веб-сервис для интернет-магазина продуктов «Славянка», предназначенный для онлайн-продажи продуктов питания, предоставления информации о магазине, его деятельности и ассортименте. Сервис позволяет пользователям просматривать каталог товаров, добавлять товары в корзину, оформлять заказы, отслеживать их статус в личном кабинете, а также общаться с администратором через встроенный чат.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Источники разработки: Среда разработки — ОС Windows 11, локальная разработка и написание кода в Visual Studio Code. Стек технологий: Vue 3, TypeScript, Fastify, Prisma, SQLite.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Требования к функциональным характеристикам:', bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    reqs = [
        'наличие каталога товаров с категориями, поиском и фильтрацией;',
        'наличие корзины для оформления заказов;',
        'регистрация и авторизация пользователей;',
        'личный кабинет пользователя с историей заказов и избранными товарами;',
        'административная панель для управления товарами, категориями, заказами, пользователями, новостями и чатами;',
        'система новостей и акций;',
        'чат для обратной связи с администратором;',
        'возможность просмотра и добавления отзывов на товары;',
        'умный адрес с автодополнением при оформлении заказа;',
        'промокоды и скидки.',
    ]
    for r in reqs:
        add_paragraph(doc, f'– {r}', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_empty_line(doc)
    add_paragraph(doc, 'Требования к надежности: защита сервера от внезапного отключения питания и перепадов напряжения. Устойчивость к атакам из сети Интернет (Helmet, rate-limit). Обеспечение сохранности персональных данных пользователей.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_paragraph(doc, 'Требования к информационной и программной совместимости: для стабильной работы необходим браузер с поддержкой актуальных версий JavaScript, HTML, CSS. Для отображения графических элементов необходимо стабильное интернет-соединение.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_empty_line(doc)
    add_empty_line(doc)
    add_paragraph(doc, 'Дата выдачи «___» _______________\t\tСрок исполнения «___» _______________', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_empty_line(doc)
    add_paragraph(doc, 'Студент\t\t\t\t\tГ.М. Жуков', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Руководитель проекта\t\t\tИ.О. Меренцева', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Старший консультант\t\t\tИ.О. Меренцева', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.save(os.path.join(OUTPUT_DIR, '2 Задание ДП.doc'))
    print('[OK] 2 Задание ДП.doc')

# ============================================================
# ГЕНЕРАЦИЯ ТЕХНИЧЕСКОГО ЗАДАНИЯ
# ============================================================
def generate_tz():
    doc = Document()
    setup_page(doc)
    set_normal_style(doc)
    add_frame(doc)
    
    add_heading_centered(doc, 'Техническое задание', size=16)
    add_heading_centered(doc, 'НА РАЗРАБОТКУ И РЕШЕНИЕ ИТ ЗАДАЧ', size=14)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Студент: Жуков Григорий Максимович', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Специальность: 09.02.07 Информационные системы и программирование', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_paragraph(doc, 'Тема: Разработка веб-сервиса для интернет-магазина продуктов «Славянка».', bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Краткая характеристика: Веб-сервис для интернет-магазина продуктов «Славянка», предназначенный для презентации товаров магазина, ознакомления пользователей с ассортиментом, акциями и новостями. Сервис позволяет пользователям просматривать каталог, оформлять заказы онлайн и отслеживать их статус.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Источники разработки: Среда разработки — ОС Windows 11, Visual Studio Code. Стек: Vue 3, TypeScript, Fastify, Prisma, SQLite.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Требования к функциональным характеристикам:', bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Информационный продукт должен предоставлять информацию о деятельности магазина. Веб-сервис должен иметь страницу каталога с категориями товаров. Должна быть реализована функциональность корзины и оформления заказов. Должна быть реализована панель администратора для управления товарами, заказами и контентом сайта.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Требования к надежности: защита сервера от внезапного отключения питания. Устойчивость к атакам из сети Интернет. Обеспечение сохранности персональных данных пользователей.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Требования к информационной и программной совместимости: для стабильной работы интерфейса необходимо наличие браузера с поддержкой актуальных версий JavaScript, HTML, CSS.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Дата выдачи «___» _______________\t\tСрок исполнения «___» _______________', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_paragraph(doc, 'Студент\t\t\t\t\tГ.М. Жуков', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Руководитель проекта\t\t\tИ.О. Меренцева', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Старший консультант\t\t\tИ.О. Меренцева', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.save(os.path.join(OUTPUT_DIR, '3 ТЗ.docx'))
    print('[OK] 3 ТЗ.docx')

# ============================================================
# ГЕНЕРАЦИЯ АННОТАЦИИ
# ============================================================
def generate_annotaciya():
    doc = Document()
    setup_page(doc)
    set_normal_style(doc)
    add_frame(doc)
    
    add_heading_centered(doc, 'Аннотация', size=16)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Целью данного дипломного проекта является разработка веб-сервиса для интернет-магазина продуктов «Славянка», предназначенного для автоматизации процесса онлайн-продажи продуктов питания, предоставления информации о магазине, его ассортименте и деятельности.')
    add_empty_line(doc)
    
    add_paragraph(doc, 'Программный продукт разработан в среде Visual Studio Code с использованием технологий Vue 3 (Composition API), TypeScript, Fastify, Prisma и SQLite.')
    add_empty_line(doc)
    
    add_paragraph(doc, 'Для достижения поставленной цели были использованы следующие методы и инструменты:')
    items = [
        'анализ предметной области и требований к разрабатываемой системе;',
        'проектирование структуры веб-сервиса и пользовательского интерфейса;',
        'разработка клиентской и серверной частей приложения;',
        'создание базы данных;',
        'разработка административной панели управления;',
        'тестирование и отладка функциональных возможностей веб-сервиса;',
        'реализация системы оформления заказов и обработки пользовательских данных.',
    ]
    for item in items:
        add_paragraph(doc, f'− {item}')
    
    add_empty_line(doc)
    add_paragraph(doc, 'В результате выполнения дипломного проекта был разработан веб-сервис «Славянка», обеспечивающий автоматизацию обработки заказов и повышение удобства взаимодействия покупателей с магазином. Разработанная система обладает современным пользовательским интерфейсом, удобной навигацией и возможностью дальнейшего расширения функциональных возможностей.')
    add_empty_line(doc)
    
    add_paragraph(doc, 'Пояснительная записка содержит 75 листов. Включает в себя: 44 рисунка, 5 таблиц, 3 приложения. Количество использованных источников – 12.')
    
    add_empty_line(doc)
    add_empty_line(doc)
    add_empty_line(doc)
    
    # English version
    add_heading_centered(doc, 'Annotation', size=16)
    add_empty_line(doc)
    add_paragraph(doc, 'The purpose of this thesis project is to develop a web service for the online grocery store "Slavyanka", designed to automate the process of online food sales, provide information about the store, its assortment and activities.')
    add_empty_line(doc)
    add_paragraph(doc, 'The software product is developed in the Visual Studio Code environment using Vue 3 (Composition API), TypeScript, Fastify, Prisma, and SQLite technologies.')
    add_empty_line(doc)
    add_paragraph(doc, 'As a result of the diploma project, a web service "Slavyanka" was developed, which provides automation of order processing and enhances the convenience of customer interaction with the store.')
    
    doc.save(os.path.join(OUTPUT_DIR, '4 Аннотация.docx'))
    print('[OK] 4 Аннотация.docx')

# ============================================================
# ГЕНЕРАЦИЯ ВЕДОМОСТИ ДОКУМЕНТОВ
# ============================================================
def generate_vedomost():
    doc = Document()
    setup_page(doc, left_margin=2, other_margin=1)
    set_normal_style(doc)
    
    add_heading_centered(doc, 'ВЕДОМОСТЬ ДОКУМЕНТОВ', size=16)
    add_empty_line(doc)
    
    # Таблица
    table = doc.add_table(rows=11, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ['№ п/п', 'Наименование документа', 'Кол-во листов']
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
    
    docs = [
        ('1', 'Пояснительная записка', '75'),
        ('2', 'Техническое задание', '2'),
        ('3', 'Аннотация', '1'),
        ('4', 'Ведомость документов', '1'),
        ('5', 'Чертежи (карта сайта, схема БД)', '2'),
        ('6', 'Листинг программы (приложение А)', '10'),
        ('7', 'Карта сайта (приложение Б)', '1'),
        ('8', 'Схема базы данных (приложение В)', '1'),
        ('9', 'Презентация', '1'),
        ('10', 'Отчёт об антиплагиате', '1'),
    ]
    
    for i, (num, name, pages) in enumerate(docs):
        set_cell_text(table.cell(i+1, 0), num)
        set_cell_text(table.cell(i+1, 1), name, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(i+1, 2), pages)
    
    doc.save(os.path.join(OUTPUT_DIR, '6 Ведомость документов.doc'))
    print('[OK] 6 Ведомость документов.doc')

# ============================================================
# ГЕНЕРАЦИЯ СПРАВКИ О ВНЕДРЕНИИ
# ============================================================
def generate_spravka():
    doc = Document()
    setup_page(doc, left_margin=3, other_margin=2)
    set_normal_style(doc)
    
    add_heading_centered(doc, 'СПРАВКА О ВНЕДРЕНИИ', size=16)
    add_empty_line(doc)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Настоящим подтверждается, что результаты дипломного проекта на тему «Разработка веб-сервиса для интернет-магазина продуктов «Славянка»», выполненного студентом группы ИС-422 Жуковым Григорием Максимовичем, были внедрены в деятельность магазина «Славянка» (д. Курилово).')
    add_empty_line(doc)
    
    add_paragraph(doc, 'Внедрение разработанного веб-сервиса позволило:')
    items = [
        'автоматизировать процесс приёма и обработки заказов;',
        'сократить время обслуживания клиентов;',
        'повысить удобство взаимодействия покупателей с магазином;',
        'организовать централизованное управление товарами и заказами.',
    ]
    for item in items:
        add_paragraph(doc, f'– {item}')
    
    add_empty_line(doc)
    add_empty_line(doc)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Директор магазина «Славянка»', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc)
    add_paragraph(doc, '«___» _______________ 2026 г.', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.save(os.path.join(OUTPUT_DIR, 'Справка о внедрении.docx'))
    print('[OK] Справка о внедрении.docx')

# ============================================================
# ГЕНЕРАЦИЯ ПРИЛОЖЕНИЯ А (ЛИСТИНГ)
# ============================================================
def generate_prilozhenie_a():
    doc = Document()
    setup_page(doc)
    set_normal_style(doc)
    add_frame(doc)
    
    add_heading_centered(doc, 'Приложение А', size=16)
    add_heading_centered(doc, 'Листинг программы', size=14)
    add_empty_line(doc)
    
    add_paragraph(doc, 'В данном приложении представлены ключевые фрагменты исходного кода веб-сервиса «Славянка». Листинги содержат основные модули клиентской и серверной частей приложения.')
    add_empty_line(doc)
    
    # А.1
    add_subheading(doc, 'А.1 Точка входа сервера (server.ts)', size=14)
    add_paragraph(doc, 'Листинг А.1 – Главный файл сервера Fastify', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import dotenv from 'dotenv';
import path from 'path';
dotenv.config({ path: path.resolve(__dirname, '../.env') });

import Fastify from 'fastify';
import cors from '@fastify/cors';
import jwt from '@fastify/jwt';
import swagger from '@fastify/swagger';
import swaggerUi from '@fastify/swagger-ui';
import multipart from '@fastify/multipart';
import helmet from '@fastify/helmet';
import rateLimit from '@fastify/rate-limit';
import { registerRoutes } from './routes';

const PORT = parseInt(process.env.PORT || '3001', 10);
const HOST = process.env.HOST || '0.0.0.0';

async function main() {
  const server = Fastify({ logger: { level: 'info' } });
  
  // Регистрация плагинов
  await server.register(cors, { origin: true, credentials: true });
  await server.register(jwt, { secret: process.env.JWT_SECRET || 'slavyanka-super-secret-key-2026' });
  await server.register(swagger, {
    openapi: {
      info: { title: 'Славянка API', description: 'API для магазина "Славянка"', version: '1.0.0' },
    },
  });
  await server.register(swaggerUi, { routePrefix: '/docs' });
  await server.register(multipart, { limits: { fileSize: 10 * 1024 * 1024 } });
  await server.register(helmet, { contentSecurityPolicy: false });
  await server.register(rateLimit, { max: 100, timeWindow: '1 minute' });
  
  // Подключение маршрутов
  await registerRoutes(server);
  server.get('/health', async () => ({ status: 'ok' }));
  
  try {
    await server.listen({ port: PORT, host: HOST });
    console.log(`Server running on http://${HOST}:${PORT}`);
  } catch (err) {
    server.log.error(err);
    process.exit(1);
  }
}
main();"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.2
    add_subheading(doc, 'А.2 Регистрация маршрутов API (routes/index.ts)', size=14)
    add_paragraph(doc, 'Листинг А.2 – Подключение всех маршрутов сервера', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { FastifyInstance } from 'fastify';
import { authRoutes } from './auth.routes';
import { productsRoutes } from './products.routes';
import { categoriesRoutes } from './categories.routes';
import { ordersRoutes } from './orders.routes';
import { usersRoutes } from './users.routes';
import { notificationsRoutes } from './notifications.routes';
import { visitsRoutes } from './visits.routes';
import { uploadRoutes } from './upload.routes';
import { chatRoutes } from './chat.routes';
import { reviewsRoutes } from './reviews.routes';
import { newsRoutes } from './news.routes';
import { promoRoutes } from './promo.routes';
import { contactMessagesRoutes } from './contact-messages.routes';

export async function registerRoutes(server: FastifyInstance) {
  server.register(authRoutes, { prefix: '/api/auth' });
  server.register(productsRoutes, { prefix: '/api/products' });
  server.register(categoriesRoutes, { prefix: '/api/categories' });
  server.register(ordersRoutes, { prefix: '/api/orders' });
  server.register(usersRoutes, { prefix: '/api/users' });
  server.register(notificationsRoutes, { prefix: '/api/notifications' });
  server.register(visitsRoutes, { prefix: '/api/visits' });
  server.register(uploadRoutes, { prefix: '/api' });
  server.register(chatRoutes, { prefix: '/api' });
  server.register(reviewsRoutes, { prefix: '/api/reviews' });
  server.register(newsRoutes, { prefix: '/api/news' });
  server.register(promoRoutes, { prefix: '/api/promo' });
  server.register(contactMessagesRoutes, { prefix: '/api' });
}"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.3
    add_subheading(doc, 'А.3 Схема базы данных (schema.prisma)', size=14)
    add_paragraph(doc, 'Листинг А.3 – Модели данных Prisma', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """generator client { provider = "prisma-client-js" }
datasource db { provider = "sqlite" url = "file:./dev.db" }

model User {
  id            String   @id @default(uuid())
  email         String   @unique
  password      String   // bcrypt hash
  name          String
  role          String   @default("customer")
  verified      Boolean  @default(false)
  createdAt     DateTime @default(now())
  updatedAt     DateTime @updatedAt
  orders        Order[]
  notifications ProductNotification[]
  chatSessions  ChatSession[]
  reviews       Review[]
  @@map("users")
}

model VerificationCode {
  id        String   @id @default(uuid())
  email     String
  code      String
  expiresAt DateTime
  createdAt DateTime @default(now())
  @@index([email, code])
  @@map("verification_codes")
}

model Category {
  id          String    @id @default(uuid())
  name        String
  slug        String    @unique
  description String?
  imageUrl    String?
  createdAt   DateTime  @default(now())
  updatedAt   DateTime  @updatedAt
  products    Product[]
  @@map("categories")
}

model Product {
  id            String   @id @default(uuid())
  name          String
  description   String?
  price         Float
  weight        String?
  imageUrl      String?
  categoryId    String
  inStock       Boolean  @default(true)
  composition   String?
  manufacturer  String?
  avgRating     Float    @default(0)
  reviewCount   Int      @default(0)
  discountPrice Float?
  isPromo       Boolean  @default(false)
  deliveryDate  DateTime?
  category      Category @relation(fields: [categoryId], references: [id])
  reviews       Review[]
  orderItems    OrderItem[]
  @@map("products")
}

model Order {
  id              String     @id @default(uuid())
  userId          String
  status          String     @default("new")
  totalAmount     Float
  customerName    String     @default("")
  customerPhone   String     @default("")
  customerAddress String     @default("")
  deliveryType    String     @default("pickup")
  deliveryTime    String?
  paymentMethod   String     @default("")
  paid            Boolean    @default(false)
  createdAt       DateTime   @default(now())
  updatedAt       DateTime   @updatedAt
  user            User       @relation(fields: [userId], references: [id])
  items           OrderItem[]
  @@map("orders")
}

model OrderItem {
  id        String   @id @default(uuid())
  orderId   String
  productId String
  quantity  Int
  price     Float
  product   Product  @relation(fields: [productId], references: [id])
  order     Order    @relation(fields: [orderId], references: [id])
  @@map("order_items")
}

model Review {
  id        String   @id @default(uuid())
  userId    String
  productId String
  rating    Int
  text      String?
  photo     String?
  createdAt DateTime @default(now())
  user      User     @relation(fields: [userId], references: [id])
  product   Product  @relation(fields: [productId], references: [id])
  @@unique([userId, productId])
  @@map("reviews")
}

model News {
  id        String   @id @default(uuid())
  title     String
  slug      String   @unique
  content   String
  imageUrl  String?
  excerpt   String?
  published Boolean  @default(false)
  createdAt DateTime @default(now())
  updatedAt DateTime @updatedAt
  @@map("news")
}

model ChatSession {
  id        String        @id @default(uuid())
  userId    String?
  guestId   String
  userName  String?
  status    String        @default("active")
  createdAt DateTime      @default(now())
  updatedAt DateTime      @updatedAt
  messages  ChatMessage[]
  user      User?         @relation(fields: [userId], references: [id])
  @@map("chat_sessions")
}

model ChatMessage {
  id         String       @id @default(uuid())
  sessionId  String
  senderType String       // 'user' или 'admin'
  content    String
  attachment String?
  isRead     Boolean      @default(false)
  createdAt  DateTime     @default(now())
  session    ChatSession  @relation(fields: [sessionId], references: [id], onDelete: Cascade)
  @@map("chat_messages")
}

model PromoCode {
  id        String   @id @default(uuid())
  email     String
  code      String   @unique
  used      Boolean  @default(false)
  createdAt DateTime @default(now())
  expiresAt DateTime
  @@map("promo_codes")
}

model ContactMessage {
  id         String    @id @default(uuid())
  name       String
  email      String
  phone      String    @default("")
  topic      String    @default("")
  message    String
  adminReply String?
  isRead     Boolean   @default(false)
  createdAt  DateTime  @default(now())
  updatedAt  DateTime  @updatedAt
  @@map("contact_messages")
}

model AuditLog {
  id        String   @id @default(uuid())
  userId    String?
  action    String
  entity    String?
  details   String?
  ipAddress String?
  createdAt DateTime @default(now())
  user      User?    @relation(fields: [userId], references: [id], onDelete: SetNull)
  @@map("audit_logs")
}"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.4
    add_subheading(doc, 'А.4 Хранилище корзины (cart.store.ts)', size=14)
    add_paragraph(doc, 'Листинг А.4 – Pinia store для управления корзиной', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { defineStore } from 'pinia';
import { ref, computed } from 'vue';

// Хранилище корзины с использованием Pinia Composition API
export const useCartStore = defineStore('cart', () => {
  const items = ref([]);

  // Вычисляемые свойства
  const totalItems = computed(() => 
    items.value.reduce((sum, item) => sum + item.quantity, 0)
  );
  const totalPrice = computed(() => 
    items.value.reduce((sum, item) => sum + item.price * item.quantity, 0)
  );

  // Добавление товара в корзину
  function addItem(product) {
    const existing = items.value.find(i => i.id === product.id);
    if (existing) {
      existing.quantity++;
    } else {
      items.value.push({ ...product, quantity: 1 });
    }
  }

  // Удаление товара из корзины
  function removeItem(productId) {
    items.value = items.value.filter(i => i.id !== productId);
  }

  // Изменение количества товара
  function updateQuantity(productId, qty) {
    const item = items.value.find(i => i.id === productId);
    if (item) {
      item.quantity = Math.max(1, qty);
    }
  }

  // Очистка корзины
  function clearCart() {
    items.value = [];
  }

  return { items, totalItems, totalPrice, addItem, removeItem, updateQuantity, clearCart };
});"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.5
    add_subheading(doc, 'А.5 Хранилище авторизации (auth.store.ts)', size=14)
    add_paragraph(doc, 'Листинг А.5 – Pinia store для аутентификации', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { defineStore } from 'pinia';
import { ref, computed } from 'vue';
import { authApi } from '../api/auth.api';

function getLocalStorageItem(key) {
  try { return localStorage.getItem(key); }
  catch { return null; }
}
function setLocalStorageItem(key, value) {
  try { localStorage.setItem(key, value); }
  catch (error) { console.error('localStorage недоступен:', error); }
}
function removeLocalStorageItem(key) {
  try { localStorage.removeItem(key); }
  catch (error) { console.error('localStorage недоступен:', error); }
}

export const useAuthStore = defineStore('auth', () => {
  const token = ref(getLocalStorageItem('token'));
  const user = ref(null);

  const isAuthenticated = computed(() => !!token.value);
  const isAdmin = computed(() => user.value?.role === 'admin');

  async function login(email, password) {
    const response = await authApi.login(email, password);
    token.value = response.token;
    user.value = response.user;
    setLocalStorageItem('token', response.token);
  }

  async function register(email, password, name) {
    const response = await authApi.register(email, password, name);
    token.value = response.token;
    user.value = response.user;
    setLocalStorageItem('token', response.token);
  }

  async function fetchUser() {
    if (!token.value) return;
    user.value = await authApi.getMe();
  }

  function logout() {
    token.value = null;
    user.value = null;
    removeLocalStorageItem('token');
    // Очистка данных чата при выходе
    ['chat-session-id','chat-guest-id','chat-messages','chat-user-hash']
      .forEach(key => removeLocalStorageItem(key));
  }

  return { token, user, isAuthenticated, isAdmin, login, register, fetchUser, logout };
});"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.6
    add_subheading(doc, 'А.6 Хранилище избранного (favorites.store.ts)', size=14)
    add_paragraph(doc, 'Листинг А.6 – Модуль избранного с localStorage', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { ref, computed, watch } from 'vue';

const GUEST_STORAGE_KEY = 'slavyanka-favorites-guest';
let storageKey = GUEST_STORAGE_KEY;

// Тип элемента избранного
export interface FavoriteItem {
  id: string;
  name: string;
  price: number;
  imageUrl?: string;
  inStock: boolean;
  avgRating?: number;
}

const favorites = ref([]);

// Загрузка избранного из localStorage
function loadFavorites() {
  try {
    const saved = localStorage.getItem(storageKey);
    favorites.value = saved ? JSON.parse(saved) : [];
  } catch (e) {
    favorites.value = [];
  }
}

// Сохранение избранного в localStorage
function saveFavorites() {
  try {
    localStorage.setItem(storageKey, JSON.stringify(favorites.value));
  } catch (e) {
    console.error('Failed to save favorites:', e);
  }
}

loadFavorites();

// Автосохранение при изменениях
watch(favorites, () => { saveFavorites(); }, { deep: true });

export function resetFavorites() { favorites.value = []; }

export function setFavoritesUserId(userId) {
  const newKey = userId ? `slavyanka-favorites-${userId}` : GUEST_STORAGE_KEY;
  if (newKey !== storageKey) {
    storageKey = newKey;
    loadFavorites();
  }
}

export function useFavorites() {
  const count = computed(() => favorites.value.length);

  function toggle(item) {
    const index = favorites.value.findIndex(f => f.id === item.id);
    if (index === -1) favorites.value.push(item);
    else favorites.value.splice(index, 1);
  }

  function isFavorite(id) { return favorites.value.some(f => f.id === id); }
  function remove(id) {
    const index = favorites.value.findIndex(f => f.id === id);
    if (index !== -1) favorites.value.splice(index, 1);
  }
  function getAll() { return [...favorites.value]; }

  return { favorites, count, toggle, isFavorite, remove, getAll };
}"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.7
    add_subheading(doc, 'А.7 API-клиент для авторизации (auth.api.ts)', size=14)
    add_paragraph(doc, 'Листинг А.7 – Клиент для взаимодействия с сервером авторизации', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import apiClient from './client';

// API для работы с аутентификацией
export const authApi = {
  // Вход в систему
  async login(email, password) {
    const response = await apiClient.post('/auth/login', { email, password });
    return response.data;
  },

  // Регистрация нового пользователя
  async register(email, password, name) {
    const response = await apiClient.post('/auth/register', { email, password, name });
    return response.data;
  },

  // Получение информации о текущем пользователе
  async getMe() {
    const response = await apiClient.get('/auth/me');
    return response.data;
  },

  // Отправка кода подтверждения email
  async sendVerification(email) {
    const response = await apiClient.post('/auth/send-verification', { email });
    return response.data;
  },

  // Проверка кода подтверждения
  async verifyCode(email, code) {
    const response = await apiClient.post('/auth/verify-code', { email, code });
    return response.data;
  },
};"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.8
    add_subheading(doc, 'А.8 API-клиент для товаров (products.api.ts)', size=14)
    add_paragraph(doc, 'Листинг А.8 – Клиент для работы с каталогом товаров', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import apiClient from './client';

export const productsApi = {
  async getAll(params = {}) {
    const response = await apiClient.get('/products', { params });
    return response.data;
  },

  async getById(id) {
    const response = await apiClient.get(`/products/${id}`);
    return response.data;
  },

  async create(data) {
    const response = await apiClient.post('/products', data);
    return response.data;
  },

  async update(id, data) {
    const response = await apiClient.put(`/products/${id}`, data);
    return response.data;
  },

  async delete(id) {
    const response = await apiClient.delete(`/products/${id}`);
    return response.data;
  },
};"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.9
    add_subheading(doc, 'А.9 Валидация форм (validation.ts)', size=14)
    add_paragraph(doc, 'Листинг А.9 – Функции валидации и защита от XSS', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """// Валидация email
export function validateEmail(email) {
  return /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/.test(email);
}

// Валидация номера телефона
export function validatePhone(phone) {
  return /^[+]?[\\d\\s()-]{10,15}$/.test(phone);
}

// Проверка заполненности поля
export function validateRequired(value) {
  return value.trim().length > 0;
}

// Санитизация ввода для защиты от XSS-атак
export function sanitizeInput(input) {
  return input
    .replace(/&/g, '&')
    .replace(/</g, '<')
    .replace(/>/g, '>')
    .replace(/"/g, '"')
    .replace(/'/g, '&#x27;');
}"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.10
    add_subheading(doc, 'А.10 Хеширование паролей (hash.ts)', size=14)
    add_paragraph(doc, 'Листинг А.10 – Утилита для хеширования паролей (bcrypt)', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import bcrypt from 'bcryptjs';

const SALT_ROUNDS = 10;

// Хеширование пароля
export async function hashPassword(password) {
  return bcrypt.hash(password, SALT_ROUNDS);
}

// Сравнение пароля с хешем
export async function comparePassword(password, hash) {
  return bcrypt.compare(password, hash);
}"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.11
    add_subheading(doc, 'А.11 Маршруты авторизации (auth.routes.ts)', size=14)
    add_paragraph(doc, 'Листинг А.11 – Обработка запросов авторизации на сервере', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { FastifyInstance } from 'fastify';
import { authController } from '../controllers/auth.controller';

export async function authRoutes(server) {
  // Регистрация нового пользователя
  server.post('/register', authController.register);
  
  // Вход в систему
  server.post('/login', authController.login);
  
  // Получение текущего пользователя (требуется JWT)
  server.get('/me', { preHandler: [server.authenticate] }, authController.getMe);
  
  // Отправка кода подтверждения email
  server.post('/send-verification', authController.sendVerification);
  
  // Проверка кода подтверждения
  server.post('/verify-code', authController.verifyCode);
}"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.12
    add_subheading(doc, 'А.12 Сервис авторизации (auth.service.ts)', size=14)
    add_paragraph(doc, 'Листинг А.12 – Бизнес-логика аутентификации', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { PrismaClient } from '@prisma/client';
import { hashPassword, comparePassword } from '../utils/hash';

const prisma = new PrismaClient();

export const authService = {
  // Регистрация
  async register({ email, password, name }) {
    const existing = await prisma.user.findUnique({ where: { email } });
    if (existing) throw new Error('Email уже используется');
    
    const hashedPassword = await hashPassword(password);
    const user = await prisma.user.create({
      data: { email, password: hashedPassword, name },
    });
    return { id: user.id, email: user.email, name: user.name, role: user.role };
  },

  // Вход
  async login({ email, password }) {
    const user = await prisma.user.findUnique({ where: { email } });
    if (!user) throw new Error('Неверный email или пароль');
    
    const valid = await comparePassword(password, user.password);
    if (!valid) throw new Error('Неверный email или пароль');
    
    return { id: user.id, email: user.email, name: user.name, role: user.role };
  },
};"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.13
    add_subheading(doc, 'А.13 Контроллер товаров (products.controller.ts)', size=14)
    add_paragraph(doc, 'Листинг А.13 – Обработчик запросов к каталогу товаров', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { productsService } from '../services/products.service';

export const productsController = {
  // Получение списка товаров с фильтрацией
  async getAll(request, reply) {
    const { category, search, sort, page, limit } = request.query;
    const products = await productsService.getAll({
      category, search, sort, page, limit,
    });
    return products;
  },

  // Получение одного товара
  async getById(request, reply) {
    const { id } = request.params;
    const product = await productsService.getById(id);
    if (!product) {
      return reply.status(404).send({ error: 'Товар не найден' });
    }
    return product;
  },

  // Создание товара (только админ)
  async create(request, reply) {
    const product = await productsService.create(request.body);
    return reply.status(201).send(product);
  },

  // Обновление товара
  async update(request, reply) {
    const { id } = request.params;
    const product = await productsService.update(id, request.body);
    return product;
  },

  // Удаление товара
  async delete(request, reply) {
    const { id } = request.params;
    await productsService.delete(id);
    return reply.status(204).send();
  },
};"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.14
    add_subheading(doc, 'А.14 Сервис заказов (orders.service.ts)', size=14)
    add_paragraph(doc, 'Листинг А.14 – Бизнес-логика обработки заказов', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { PrismaClient } from '@prisma/client';

const prisma = new PrismaClient();

export const ordersService = {
  async getAll({ userId, status, page = 1, limit = 20 }) {
    const where = {};
    if (userId) where.userId = userId;
    if (status) where.status = status;

    const [items, total] = await Promise.all([
      prisma.order.findMany({
        where,
        include: { items: { include: { product: true } } },
        skip: (page - 1) * limit,
        take: limit,
        orderBy: { createdAt: 'desc' },
      }),
      prisma.order.count({ where }),
    ]);
    return { items, total, page, limit };
  },

  async getById(id) {
    return prisma.order.findUnique({
      where: { id },
      include: { items: { include: { product: true } } },
    });
  },

  async create(data) {
    const { userId, items, customerName, customerPhone, deliveryAddress } = data;
    const totalAmount = items.reduce(
      (sum, item) => sum + item.price * item.quantity, 0
    );

    return prisma.order.create({
      data: {
        userId,
        totalAmount,
        customerName,
        customerPhone,
        deliveryAddress,
        items: {
          create: items.map(item => ({
            productId: item.productId,
            quantity: item.quantity,
            price: item.price,
          })),
        },
      },
      include: { items: true },
    });
  },

  async updateStatus(id, status) {
    return prisma.order.update({ where: { id }, data: { status } });
  },
};"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc); doc.add_page_break()
    
    # А.15
    add_subheading(doc, 'А.15 Контроллер чата (chat.controller.ts)', size=14)
    add_paragraph(doc, 'Листинг А.15 – Обработка WebSocket-чата поддержки', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    code = """import { PrismaClient } from '@prisma/client';

const prisma = new PrismaClient();

export const chatController = {
  // Создание новой сессии чата
  async createSession(request, reply) {
    const { userId, guestId, userName } = request.body;
    const session = await prisma.chatSession.create({
      data: { userId, guestId, userName },
    });
    return session;
  },

  // Получение сообщений сессии
  async getMessages(request, reply) {
    const { sessionId } = request.params;
    const messages = await prisma.chatMessage.findMany({
      where: { sessionId },
      orderBy: { createdAt: 'asc' },
    });
    return messages;
  },

  // Отправка сообщения
  async sendMessage(request, reply) {
    const { sessionId, senderType, content } = request.body;
    const message = await prisma.chatMessage.create({
      data: { sessionId, senderType, content },
    });
    return message;
  },

  // Получение всех активных сессий (для админа)
  async getActiveSessions(request, reply) {
    const sessions = await prisma.chatSession.findMany({
      where: { status: 'active' },
      include: { messages: { take: 1, orderBy: { createdAt: 'desc' } } },
      orderBy: { updatedAt: 'desc' },
    });
    return sessions;
  },
};"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    add_empty_line(doc)
    
    add_empty_line(doc)
    add_paragraph(doc, 'Полный текст программы веб-сервиса представлен в электронном виде и прилагается к дипломному проекту на CD-диске.', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.save(os.path.join(OUTPUT_DIR, 'Приложения', 'Приложение А (Листинг).docx'))
    print('[OK] Приложение А (Листинг).docx')

# ============================================================
# ГЕНЕРАЦИЯ ПРИЛОЖЕНИЯ Б (КАРТА САЙТА)
# ============================================================
def generate_prilozhenie_b():
    doc = Document()
    setup_page(doc)
    set_normal_style(doc)
    add_frame(doc)
    
    add_heading_centered(doc, 'Приложение Б', size=16)
    add_heading_centered(doc, 'Карта сайта', size=14)
    add_empty_line(doc)
    
    add_paragraph(doc, 'На рисунке Б.1 представлена карта сайта веб-сервиса интернет-магазина продуктов «Славянка».', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    # Текстовая карта сайта
    add_heading_section(doc, 'Структура веб-сервиса «Славянка»:', size=14)
    add_empty_line(doc)
    
    map_lines = [
        'Главная страница (/)',
        '  ├── Баннер с акциями',
        '  ├── Преимущества магазина',
        '  ├── Популярные товары',
        '  ├── Отзывы клиентов',
        '  └── Карта проезда',
        'Каталог товаров (/catalog)',
        '  ├── Фильтрация по категориям',
        '  ├── Поиск по названию',
        '  ├── Сортировка (цена, алфавит)',
        '  └── Модальное окно товара',
        'Акции и скидки (/promo)',
        '  └── Товары со скидками / промокоды',
        'Новости (/news)',
        '  └── Статья новости (/news/:slug)',
        'О магазине (/about)',
        'Контакты (/contacts)',
        '  └── Форма обратной связи',
        'Корзина (/cart)',
        '  └── Оформление заказа (требуется авторизация)',
        'Личный кабинет (/profile)',
        '  ├── Мои данные',
        '  ├── История заказов',
        '  ├── Избранное',
        '  └── Мои отзывы',
        'Вход (/login)',
        'Регистрация (/register)',
        'Политика конфиденциальности (/privacy)',
        'Административная панель (/admin)',
        '  ├── Дашборд (/admin/dashboard)',
        '  ├── Товары (/admin/products)',
        '  ├── Категории (/admin/categories)',
        '  ├── Заказы (/admin/orders)',
        '  ├── Пользователи (/admin/users)',
        '  ├── Чаты (/admin/chats)',
        '  ├── Новости (/admin/news)',
        '  └── Сообщения (/admin/messages)',
    ]
    
    for line in map_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
    
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок Б.1 – Карта сайта веб-сервиса «Славянка»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.save(os.path.join(OUTPUT_DIR, 'Приложения', 'Приложение Б (Карта сайта).docx'))
    print('[OK] Приложение Б (Карта сайта).docx')

# ============================================================
# ГЕНЕРАЦИЯ ПРИЛОЖЕНИЯ В (БАЗА ДАННЫХ)
# ============================================================
def generate_prilozhenie_v():
    doc = Document()
    setup_page(doc)
    set_normal_style(doc)
    add_frame(doc)
    
    add_heading_centered(doc, 'Приложение В', size=16)
    add_heading_centered(doc, 'Схема базы данных', size=14)
    add_empty_line(doc)
    
    add_paragraph(doc, 'На рисунке В.1 представлена схема базы данных веб-сервиса интернет-магазина продуктов «Славянка». База данных реализована на СУБД SQLite с использованием ORM Prisma.')
    add_empty_line(doc)
    
    # Описание моделей в виде таблицы
    add_table_caption(doc, 'Таблица В.1 – Модели базы данных')
    
    table = doc.add_table(rows=10, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ['Модель', 'Описание', 'Ключевые поля', 'Связи']
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=11)
    
    models = [
        ('User', 'Пользователи', 'id, email, password, name, role, verified', 'orders, reviews, chatSessions'),
        ('Category', 'Категории товаров', 'id, name, slug', 'products'),
        ('Product', 'Товары', 'id, name, price, categoryId, avgRating, discountPrice', 'category, reviews, orderItems'),
        ('Order', 'Заказы', 'id, userId, status, totalAmount, customerName, deliveryAddress', 'user, items'),
        ('OrderItem', 'Позиции заказа', 'id, orderId, productId, quantity, price', 'order, product'),
        ('Review', 'Отзывы', 'id, userId, productId, rating, text', 'user, product'),
        ('News', 'Новости', 'id, title, slug, content, published', '–'),
        ('ChatSession', 'Сессии чата', 'id, userId, guestId, status', 'user, messages'),
        ('ChatMessage', 'Сообщения чата', 'id, sessionId, senderType, content', 'session'),
    ]
    
    for i, (model, desc, fields, relations) in enumerate(models):
        set_cell_text(table.cell(i+1, 0), model, size=10)
        set_cell_text(table.cell(i+1, 1), desc, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(i+1, 2), fields, size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(i+1, 3), relations, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Связи между моделями:', bold=True)
    add_paragraph(doc, '– User 1:N Order (один пользователь → много заказов)')
    add_paragraph(doc, '– User 1:N Review (один пользователь → много отзывов)')
    add_paragraph(doc, '– Category 1:N Product (одна категория → много товаров)')
    add_paragraph(doc, '– Order 1:N OrderItem (один заказ → много позиций)')
    add_paragraph(doc, '– Product 1:N OrderItem (один товар → много позиций в заказах)')
    add_paragraph(doc, '– Product 1:N Review (один товар → много отзывов)')
    add_paragraph(doc, '– ChatSession 1:N ChatMessage (одна сессия → много сообщений)')
    
    add_empty_line(doc)
    add_paragraph(doc, 'Дополнительные индексы созданы для полей email (User), slug (Category, News), а также для userId в таблицах заказов и отзывов для ускорения поиска.', bold=False)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Рисунок В.1 – Схема базы данных веб-сервиса «Славянка»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.save(os.path.join(OUTPUT_DIR, 'Приложения', 'Приложение В (База Данных).docx'))
    print('[OK] Приложение В (База Данных).docx')

# ============================================================
# ГЕНЕРАЦИЯ ПОЯСНИТЕЛЬНОЙ ЗАПИСКИ (5 ПЗ)
# ============================================================
def generate_pz():
    doc = Document()
    setup_page(doc)
    set_normal_style(doc)
    add_frame(doc)
    
    # ============================================================
    # СОДЕРЖАНИЕ
    # ============================================================
    add_heading_centered(doc, 'СОДЕРЖАНИЕ', size=16)
    
    toc = [
        'Введение', '3',
        '1 Общая часть', '5',
        '  1.1 Постановка задачи', '5',
        '    1.1.1 Назначение задачи', '5',
        '    1.1.2 Анализ предметной области', '6',
        '    1.1.3 Требования к программе', '8',
        '    1.1.4 Требования к составу и параметрам технических средств', '9',
        '  1.2 Обоснование выбора технологий', '10',
        '    1.2.1 Выбор технологий клиентской части', '10',
        '    1.2.2 Выбор технологий серверной части', '14',
        '    1.2.3 Обоснование выбора системы управления базами данных', '16',
        '2 Специальная часть', '20',
        '  2.1 Архитектура приложения', '20',
        '    2.1.1 Структура веб-сервиса', '20',
        '    2.1.2 Схема базы данных', '23',
        '    2.1.3 Алгоритмы обработки информации', '25',
        '    2.1.4 Обеспечение безопасности', '28',
        '  2.2 Описание работы сайта', '30',
        '  2.3 Инструкция работы с сайтом', '34',
        '    2.3.1 Общие сведения', '34',
        '    2.3.2 Вызов и загрузка', '39',
        '    2.3.3 Входные данные', '40',
        '    2.3.4 Выходные данные', '40',
        '    2.3.5 Сообщения программы', '41',
        '  2.4 Описание процесса отладки программы', '42',
        '    2.4.1 Методы отладки', '42',
        '    2.4.2 Средства тестирования', '43',
        '    2.4.3 Контрольный пример', '44',
        '    2.4.4 Тестирование функциональных модулей', '46',
        '  2.5 Развёртывание и сопровождение', '47',
        '3 Экономическая часть', '49',
        'Заключение', '52',
        'Список литературы', '54',
    ]
    
    for i in range(0, len(toc), 2):
        p = doc.add_paragraph()
        if toc[i].startswith('  '):
            p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(toc[i].strip())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # ============================================================
    # ВВЕДЕНИЕ
    # ============================================================
    add_heading_centered(doc, 'ВВЕДЕНИЕ', size=16)
    
    add_paragraph(doc, 'Современное общество характеризуется высоким уровнем развития информационных технологий, которые активно внедряются практически во все сферы жизнедеятельности человека. Информация стала одним из важнейших ресурсов, а скорость её получения и обработки напрямую влияет на эффективность работы организаций и качество предоставляемых услуг. Интернет превратился в универсальную цифровую среду, объединяющую бизнес, пользователей и различные сервисы.')
    
    add_paragraph(doc, 'Информационные технологии и веб-сервисы играют важную роль в современном обществе, предоставляя множество возможностей для автоматизации процессов и повышения качества обслуживания. Всемирная сеть обеспечивает быстрый обмен информацией, упрощает взаимодействие между людьми и делает многие услуги более доступными. Использование веб-технологий позволяет организациям эффективно представлять свои товары и услуги, поддерживать обратную связь с клиентами и оптимизировать внутренние процессы.')
    
    add_paragraph(doc, 'Сфера розничной торговли продуктами питания является одной из ключевых в экономике, где качество обслуживания играет особенно важную роль. Конкуренция между продуктовыми магазинами постоянно растёт, поэтому организациям приходится уделять внимание не только качеству товаров, но и удобству взаимодействия с покупателями. Потенциальный покупатель всё чаще знакомится с ассортиментом магазина именно через интернет. Посетитель изучает товары, цены, акции, информацию о доставке и формирует решение о покупке ещё до посещения магазина.')
    
    add_paragraph(doc, 'Продуктовый магазин «Славянка» уже несколько лет работает в деревне Курилово, обеспечивая жителей свежими и качественными продуктами. Магазин предлагает широкий ассортимент товаров — от свежей выпечки и молочных продуктов до бытовых товаров. Разработка веб-сервиса для онлайн-продажи продуктов позволит расширить клиентскую базу, автоматизировать процесс приёма заказов и повысить качество обслуживания.')
    
    add_paragraph(doc, 'Актуальность данной темы обусловлена ростом спроса на онлайн-покупки продуктов среди жителей сельской местности и отсутствием у магазина «Славянка» автоматизированной системы приёма заказов. Разработка веб-сервиса позволит автоматизировать взаимодействие между магазином и покупателями, уменьшить нагрузку на сотрудников и повысить эффективность работы.')
    
    add_paragraph(doc, 'Цель дипломной работы заключается в разработке веб-сервиса для интернет-магазина продуктов «Славянка», обеспечивающего возможность онлайн-просмотра каталога товаров, оформления заказов и управления ими.')
    
    add_paragraph(doc, 'Для достижения поставленной цели необходимо решить следующие задачи:')
    tasks = [
        'выполнить анализ предметной области, определить требования к веб-сервису и обосновать выбор технологий разработки;',
        'разработать архитектуру приложения, реализовать клиентскую и серверную части, описать алгоритмы работы программного продукта;',
        'провести тестирование системы, подготовить руководство пользователя и выполнить экономическое обоснование разработки.',
    ]
    for t in tasks:
        add_paragraph(doc, f'– {t}')
    
    add_paragraph(doc, 'Объектом дипломной работы является деятельность магазина продуктов «Славянка», связанная с процессом обслуживания покупателей и организацией продаж.')
    add_paragraph(doc, 'Предметом дипломной работы выступает процесс разработки веб-сервиса для онлайн-продажи продуктов питания.')
    
    add_paragraph(doc, 'При выполнении работы использовались методы анализа предметной области, проектирования информационных систем, сравнительного анализа технологий разработки, моделирования структуры базы данных, проектирования пользовательского интерфейса, тестирования и практической реализации программного продукта.')
    
    add_paragraph(doc, 'Практическая значимость работы заключается в возможности внедрения разработанного веб-сервиса в деятельность магазина «Славянка». Использование системы позволит упростить работу сотрудников, повысить удобство взаимодействия с покупателями и улучшить организацию процесса продаж.')
    
    # ============================================================
    # 1 ОБЩАЯ ЧАСТЬ
    # ============================================================
    add_heading_centered(doc, '1 ОБЩАЯ ЧАСТЬ', size=16)
    
    add_subheading(doc, '1.1 Постановка задачи', size=15)
    
    add_subsubheading(doc, '1.1.1 Назначение задачи', size=14)
    
    add_paragraph(doc, 'От предприятия была поставлена задача разработки веб-сервиса с удобным пользовательским интерфейсом для онлайн-продажи продуктов питания. Основной целью является повышение качества обслуживания клиентов магазина за счёт автоматизации процесса оформления заказов.')
    
    add_paragraph(doc, 'В процессе работы магазина ежедневно выполняется ряд действий, связанных с обработкой обращений покупателей. Посетители уточняют информацию о товарах, наличии, ценах, делают заказы по телефону. Значительная часть данных процессов при отсутствии автоматизированной системы выполняется вручную, что увеличивает нагрузку на сотрудников.')
    
    add_paragraph(doc, 'Для повышения эффективности работы требуется автоматизация следующих процессов:')
    auto_tasks = [
        'просмотр каталога товаров с поиском и фильтрацией;',
        'добавление товаров в корзину и оформление заказа;',
        'регистрация и авторизация пользователей;',
        'просмотр истории заказов в личном кабинете;',
        'управление товарами, категориями и заказами через административную панель.',
    ]
    for t in auto_tasks:
        add_paragraph(doc, f'– {t}')
    
    add_paragraph(doc, 'Назначение разработки заключается в создании веб-сервиса, обеспечивающего автоматизацию процессов онлайн-продажи и предоставляющего пользователям удобный инструмент взаимодействия с магазином. В результате использования разработанной системы предполагается снижение нагрузки на сотрудников, уменьшение количества ошибок при обработке данных и повышение удобства обслуживания покупателей.')
    
    add_subsubheading(doc, '1.1.2 Анализ предметной области', size=14)
    
    add_paragraph(doc, 'Сфера розничной торговли продуктами питания относится к числу активно развивающихся направлений деятельности, в которых качество обслуживания и скорость взаимодействия с покупателями оказывают существенное влияние на эффективность работы организации. Современные продуктовые магазины выполняют не только функцию продажи товаров, но и предоставляют дополнительные сервисы: онлайн-заказы, доставку, программы лояльности.')
    
    add_paragraph(doc, 'Магазин «Славянка» расположен в деревне Курилово и обслуживает местное население, предлагая широкий ассортимент продуктов питания и бытовых товаров. Основными покупателями являются жители деревни и окрестных населённых пунктов.')
    
    add_paragraph(doc, 'Наиболее распространённым способом оформления заказа в настоящее время является телефонное взаимодействие между покупателем и сотрудником магазина. Во время общения продавец уточняет необходимые данные: список товаров, количество, контактную информацию. При увеличении количества обращений использование исключительно ручного способа обработки информации может приводить к различным сложностям. Среди основных проблем можно выделить:')
    problems = [
        'ошибки при записи данных;',
        'потерю информации о заказах;',
        'длительную обработку заявок;',
        'увеличение нагрузки на сотрудников;',
        'невозможность обработки заказов в нерабочее время.',
    ]
    for p in problems:
        add_paragraph(doc, f'– {p}')
    
    add_paragraph(doc, 'Для сравнения были проанализированы существующие решения для автоматизации продаж продуктов питания:')
    
    add_table_caption(doc, 'Таблица 1 – Сравнительный анализ аналогов')
    
    table1 = doc.add_table(rows=5, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = 'Table Grid'
    
    t1_headers = ['Критерий', 'Instamart', 'СберМаркет', 'Разрабатываемый сервис']
    for i, h in enumerate(t1_headers):
        set_cell_text(table1.cell(0, i), h, bold=True, size=10)
    
    t1_data = [
        ('Целевая аудитория', 'Крупные города', 'Крупные города', 'Сельская местность'),
        ('Стоимость внедрения', 'Высокая', 'Высокая', 'Низкая'),
        ('Адаптация под конкретный магазин', 'Отсутствует', 'Отсутствует', 'Полная'),
        ('Интеграция с кассой магазина', 'Отсутствует', 'Отсутствует', 'Возможна'),
    ]
    for i, (c1, c2, c3, c4) in enumerate(t1_data):
        set_cell_text(table1.cell(i+1, 0), c1, size=10)
        set_cell_text(table1.cell(i+1, 1), c2, size=10)
        set_cell_text(table1.cell(i+1, 2), c3, size=10)
        set_cell_text(table1.cell(i+1, 3), c4, size=10)
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Проведённый анализ показал, что существующие решения ориентированы на крупные города и имеют высокую стоимость внедрения, что делает их непригодными для использования в условиях небольшого сельского магазина. Разрабатываемый веб-сервис должен учитывать специфику работы магазина «Славянка» и обеспечивать:')
    spec_reqs = [
        'просмотр каталога товаров с категориями;',
        'корзину для оформления заказов;',
        'регистрацию и авторизацию пользователей;',
        'личный кабинет с историей заказов;',
        'административную панель для управления;',
        'чат для обратной связи;',
        'систему отзывов;',
        'промокоды и скидки;',
        'умный адрес с автодополнением.',
    ]
    for r in spec_reqs:
        add_paragraph(doc, f'– {r}')
    
    add_subsubheading(doc, '1.1.3 Требования к программе', size=14)
    
    add_paragraph(doc, 'По техническому заданию необходимо разработать веб-сервис, предоставляющий возможность приобретения продуктов питания через интернет. Система должна обеспечивать стабильную работу основных функций независимо от количества обращений пользователей и используемых устройств.')
    
    add_paragraph(doc, 'Функциональные возможности для различных категорий пользователей:')
    
    add_paragraph(doc, 'Для незарегистрированного пользователя:', bold=True)
    guest_reqs = [
        'просмотр каталога товаров;',
        'поиск и фильтрация товаров;',
        'просмотр информации о товаре;',
        'добавление товаров в корзину;',
        'оформление заказа (после авторизации);',
        'просмотр новостей и акций.',
    ]
    for r in guest_reqs:
        add_paragraph(doc, f'– {r}')
    
    add_paragraph(doc, 'Для зарегистрированного пользователя:', bold=True)
    user_reqs = [
        'все возможности незарегистрированного пользователя;',
        'оформление и отслеживание заказов;',
        'история заказов в личном кабинете;',
        'избранные товары;',
        'написание отзывов;',
        'чат с администратором;',
        'использование промокодов.',
    ]
    for r in user_reqs:
        add_paragraph(doc, f'– {r}')
    
    add_paragraph(doc, 'Для администратора:', bold=True)
    admin_reqs = [
        'управление товарами (добавление, редактирование, удаление);',
        'управление категориями;',
        'просмотр и изменение статусов заказов;',
        'управление пользователями;',
        'управление новостями;',
        'чат с клиентами;',
        'просмотр и ответ на сообщения из формы обратной связи.',
    ]
    for r in admin_reqs:
        add_paragraph(doc, f'– {r}')
    
    add_paragraph(doc, 'Требования к интерфейсу: понятная структура страниц, логичное расположение элементов, минимальное количество действий для оформления заказа, единый стиль оформления, адаптивность под различные устройства, удобная навигация.')
    
    add_paragraph(doc, 'Требования к производительности: стабильная работа без критических ошибок, корректная обработка запросов, быстрое отображение страниц, сохранность пользовательских данных, корректная работа на различных устройствах.')
    
    add_subsubheading(doc, '1.1.4 Требования к составу и параметрам технических средств', size=14)
    
    add_paragraph(doc, 'Программное обеспечение клиентской части должно удовлетворять требованиям в соответствии с таблицей 2.')
    
    add_table_caption(doc, 'Таблица 2 – Требования к программному обеспечению')
    
    table2 = doc.add_table(rows=3, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'
    
    set_cell_text(table2.cell(0, 0), 'Название', bold=True)
    set_cell_text(table2.cell(0, 1), 'Минимальные системные требования', bold=True)
    set_cell_text(table2.cell(1, 0), 'Операционная Система')
    set_cell_text(table2.cell(1, 1), 'Windows, Linux, macOS, Android, iOS')
    set_cell_text(table2.cell(2, 0), 'Веб-браузер')
    set_cell_text(table2.cell(2, 1), 'Google Chrome 110+, Mozilla Firefox 110+, Opera 95+, Microsoft Edge 110+, Safari 16+')
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Требования к аппаратному обеспечению представлены в соответствии с таблицей 3.')
    
    add_table_caption(doc, 'Таблица 3 – Требования к аппаратному обеспечению')
    
    table3 = doc.add_table(rows=5, cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Table Grid'
    
    set_cell_text(table3.cell(0, 0), 'Название', bold=True)
    set_cell_text(table3.cell(0, 1), 'Минимальные системные требования', bold=True)
    
    hw_reqs = [
        ('Видеокарта', 'Графическое устройство DirectX 12 с драйвером WDDM 2.7'),
        ('Процессор', '32-разрядный (x86) процессор с тактовой частотой 2 ГГц'),
        ('ОЗУ', '4 ГБ'),
        ('Интернет-соединение', '2 Мбит/с'),
    ]
    for i, (name, desc) in enumerate(hw_reqs):
        set_cell_text(table3.cell(i+1, 0), name)
        set_cell_text(table3.cell(i+1, 1), desc)
    
    doc.add_page_break()
    
    # 1.2
    add_subheading(doc, '1.2 Обоснование выбора технологий', size=15)
    
    add_subsubheading(doc, '1.2.1 Выбор технологий клиентской части', size=14)
    
    add_paragraph(doc, 'Разработка современных веб-приложений предполагает использование большого количества технологий, каждая из которых выполняет определённую задачу. От выбранных средств разработки зависит удобство использования сайта, скорость его работы, производительность системы и возможность дальнейшего расширения функциональности.')
    
    add_paragraph(doc, 'Клиентская часть веб-приложения отвечает за взаимодействие пользователя с системой и обеспечивает отображение интерфейса в браузере. Именно клиентская часть осуществляет вывод информации, обработку действий пользователя, отправку запросов на сервер и обновление содержимого страниц.')
    
    add_paragraph(doc, 'Для реализации клиентской части был выбран фреймворк Vue.js версии 3 с использованием Composition API и языка TypeScript. Vue.js является прогрессивным JavaScript-фреймворком, предназначенным для разработки пользовательских интерфейсов и одностраничных приложений. Основной особенностью данной технологии является использование компонентной архитектуры, которая позволяет разбивать интерфейс на независимые, переиспользуемые части.')
    
    add_paragraph(doc, 'К преимуществам Vue.js 3 относятся: высокая скорость работы, компонентная архитектура, реактивное обновление данных, удобная интеграция с библиотеками, поддержка адаптивного интерфейса и удобство сопровождения проекта. Для управления состоянием приложения используется Pinia — библиотека для централизованного управления состоянием, обеспечивающая реактивность данных. Для маршрутизации используется Vue Router — официальная библиотека для навигации между страницами без перезагрузки.')
    
    add_paragraph(doc, 'Для сборки и запуска клиентской части используется Vite — современный сборщик, обеспечивающий высокую скорость разработки за счёт технологии HMR (Hot Module Replacement).')
    
    add_paragraph(doc, 'Серверная часть является ключевым компонентом веб-приложения. Обработка данных, бизнес-логика и взаимодействие с базой данных выполняются именно на сервере. Для реализации серверной части был выбран фреймворк Fastify — высокопроизводительный веб-фреймворк для Node.js, написанный на TypeScript. Fastify предоставляет мощный механизм плагинов, встроенную валидацию схем, поддержку TypeScript «из коробки» и высокую скорость обработки запросов.')
    
    add_paragraph(doc, 'Для работы с базой данных используется Prisma — современная ORM (Object-Relational Mapping) для Node.js и TypeScript. Prisma обеспечивает автоматическую генерацию типов на основе схемы данных, безопасные типизированные запросы, простую систему миграций и удобный интерфейс для работы с данными.')
    
    add_paragraph(doc, 'В качестве базы данных используется SQLite — встроенная реляционная база данных, не требующая установки отдельного сервера. Данные хранятся в одном файле, что упрощает развёртывание и резервное копирование. При необходимости через Prisma можно легко переключиться на другую СУБД (PostgreSQL, MySQL).')
    
    add_paragraph(doc, 'Дополнительно в проекте используются: TypeScript для строгой типизации как клиентской, так и серверной части; HTML5 и CSS3 для разметки и стилизации; библиотека Axios для HTTP-запросов; JWT для аутентификации; bcrypt для хеширования паролей; Helmet для защиты HTTP-заголовков; rate-limit для ограничения количества запросов.')
    
    add_paragraph(doc, 'Таким образом, для разработки веб-сервиса «Славянка» был выбран следующий стек технологий: Vue 3 (Composition API), TypeScript, Pinia, Vue Router, Vite, Fastify, Prisma, SQLite.')
    
    add_empty_line(doc)
    
    # Схема архитектуры
    add_paragraph(doc, 'Архитектура взаимодействия компонентов системы представлена на рисунке 1.', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 1 – Схема архитектуры веб-сервиса «Славянка»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'Архитектура построена по принципу разделения клиентской и серверной частей. Клиентская часть (Vue 3) взаимодействует с серверной частью (Fastify) через REST API. Сервер обрабатывает запросы, выполняет бизнес-логику и взаимодействует с базой данных через Prisma ORM. Аутентификация реализована с использованием JWT-токенов.')
    
    add_subsubheading(doc, '1.2.2 Выбор технологий серверной части', size=14)
    
    add_paragraph(doc, 'Для реализации серверной части был выбран фреймворк Fastify. Данный выбор обоснован несколькими факторами. Fastify представляет собой высокопроизводительный веб-фреймфорк для Node.js, написанный на TypeScript. Одним из главных преимуществ Fastify является его скорость обработки запросов, которая существенно превышает показатели многих аналогичных решений. Встроенная система валидации схем на основе JSON Schema позволяет автоматически проверять корректность данных, поступающих от клиентской части, что снижает количество ошибок и повышает безопасность приложения.')
    
    add_paragraph(doc, 'Важным фактором при выборе Fastify стала его модульная архитектура, основанная на системе плагинов. Данный подход позволяет подключать только необходимые компоненты: CORS для кросс-доменных запросов, JWT для аутентификации, Swagger для документирования API, multipart для загрузки файлов, Helmet для защиты HTTP-заголовков и rate-limit для ограничения частоты запросов. Каждый плагин регистрируется независимо, что делает архитектуру гибкой и расширяемой.')
    
    add_paragraph(doc, 'Для работы с базой данных в проекте используется Prisma — современная объектно-реляционная проекция (ORM) для Node.js и TypeScript. Prisma обеспечивает автоматическую генерацию типов данных на основе описания схемы, что гарантирует типобезопасность запросов к базе данных. Миграции в Prisma выполняются автоматически на основе изменений схемы, что упрощает процесс обновления структуры базы данных при добавлении новых функций.')
    
    add_paragraph(doc, 'Дополнительные библиотеки, используемые на серверной стороне: bcryptjs для хеширования паролей пользователей, Nodemailer и EmailJS для отправки электронных писем (подтверждение регистрации, уведомления о статусе заказа), dotenv для управления переменными окружения.')
    
    add_paragraph(doc, 'Варианты альтернативных решений для серверной части были рассмотрены, но отклонены по следующим причинам. Express.js, являясь более популярным фреймворком, имеет меньшую производительность и не предоставляет встроенной валидации схем. NestJS, хотя и обладает более структурированной архитектурой, избыточен для данного проекта и увеличивает порог входа. Fastify является оптимальным компромиссом между производительностью, функциональностью и простотой использования.')
    
    add_empty_line(doc)
    
    add_subsubheading(doc, '1.2.3 Обоснование выбора системы управления базами данных', size=14)
    
    add_paragraph(doc, 'В качестве системы управления базами данных для проекта была выбрана SQLite. Данная СУБД представляет собой встраиваемую реляционную базу данных, которая не требует установки отдельного серверного процесса. Все данные хранятся в одном файле, что значительно упрощает развёртывание, резервное копирование и перенос проекта между различными компьютерами.')
    
    add_paragraph(doc, 'Выбор SQLite обоснован спецификой проекта. Интернет-магазин «Славянка» ориентирован на небольшое количество одновременных пользователей, поэтому использование легковесной встраиваемой базы данных является оправданным. Отсутствие необходимости в настройке и администрировании отдельного сервера баз данных снижает требования к инфраструктуре и упрощает процесс эксплуатации.')
    
    add_paragraph(doc, 'Следует отметить, что использование Prisma ORM обеспечивает высокий уровень абстракции, позволяющий при необходимости сменить СУБД без значительных изменений в коде приложения. Например, при расширении проекта можно переключиться на PostgreSQL или MySQL, изменив всего одну строку в файле конфигурации Prisma. Это даёт проекту запас для дальнейшего масштабирования без необходимости полной переработки кода.')
    
    add_paragraph(doc, 'Таким образом, выбор SQLite на начальном этапе разработки является прагматичным решением, позволяющим сосредоточиться на реализации функциональности без излишних затрат времени и ресурсов на администрирование базы данных.')
    
    add_empty_line(doc)
    
    # Конец раздела 1.2
    doc.add_page_break()
    
    # ============================================================
    # 2 СПЕЦИАЛЬНАЯ ЧАСТЬ
    # ============================================================
    add_heading_centered(doc, '2 СПЕЦИАЛЬНАЯ ЧАСТЬ', size=16)
    
    add_subheading(doc, '2.1 Архитектура приложения', size=15)
    
    add_subsubheading(doc, '2.1.1 Структура веб-сервиса', size=14)
    
    add_paragraph(doc, 'Веб-сервис «Славянка» построен по архитектуре Monorepo с использованием Turbo для управления проектом. Весь код разделён на три основных пакета: серверная часть (apps/backend), клиентская часть (apps/frontend) и общие типы (packages/shared-types).')
    
    add_paragraph(doc, 'Клиентская часть (Vue 3) организована по компонентному принципу и включает следующие основные модули:')
    frontend_modules = [
        'components/layout — базовые компоненты макета (AppHeader, SearchOverlay, ScrollToTop);',
        'components/product — компоненты товаров (ProductCard, ProductModal, ProductReviews);',
        'components/order — компоненты заказа (AddressAutocomplete);',
        'components/chat — компоненты чата (ChatWidget);',
        'components/promo — компоненты акций (PromoCodeBanner);',
        'components/ui — переиспользуемые UI-компоненты;',
        'views — страницы приложения (HomePage, CatalogPage, CartPage, ProfilePage и др.);',
        'stores — хранилища состояния Pinia (cart, auth, favorites, chat);',
        'api — клиенты для взаимодействия с REST API;',
        'router — конфигурация маршрутизации;',
        'styles — глобальные стили (global.css, animations.css);',
        'utils — вспомогательные функции (validation).',
    ]
    for m in frontend_modules:
        add_paragraph(doc, f'– {m}')
    
    add_paragraph(doc, 'Серверная часть (Fastify) организована по принципу MVC и включает:')
    backend_modules = [
        'controllers — обработчики запросов;',
        'routes — определение маршрутов;',
        'services — бизнес-логика;',
        'middleware — промежуточные обработчики (JWT, проверка ролей);',
        'utils — вспомогательные функции (hash, email);',
        'prisma — схема данных и миграции.',
    ]
    for m in backend_modules:
        add_paragraph(doc, f'– {m}')
    
    add_paragraph(doc, 'Маршрутизация клиентской части реализована с помощью Vue Router. Основные маршруты:')
    routes = [
        'Главная страница — /',
        'Каталог товаров — /catalog',
        'Акции — /promo',
        'Новости — /news, /news/:slug',
        'О магазине — /about',
        'Контакты — /contacts',
        'Корзина — /cart',
        'Личный кабинет — /profile (требуется авторизация)',
        'Вход — /login, Регистрация — /register',
        'Административная панель — /admin/* (требуются права администратора)',
    ]
    for r in routes:
        add_paragraph(doc, f'– {r}')
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Прототипы страниц веб-сервиса представлены в соответствии с рисунками 2-13.', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    for i in range(2, 14):
        add_paragraph(doc, f'Рисунок {i} – Прототип страницы веб-сервиса', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_empty_line(doc)
    
    add_subsubheading(doc, '2.1.2 Схема базы данных', size=14)
    
    add_paragraph(doc, 'База данных веб-сервиса реализована с использованием Prisma ORM и СУБД SQLite. Схема включает 12 моделей, описанных в файле schema.prisma.')
    
    add_paragraph(doc, 'Перечень моделей базы данных:')
    
    db_models = [
        ('User', 'пользователи (email, пароль (bcrypt), имя, роль, статус верификации)'),
        ('Category', 'категории товаров (название, slug)'),
        ('Product', 'товары (название, цена, вес, описание, состав, БЖУ, рейтинг, скидка)'),
        ('Order', 'заказы (статус, сумма, имя клиента, телефон, адрес, способ оплаты)'),
        ('OrderItem', 'позиции заказа (товар, количество, цена)'),
        ('Review', 'отзывы (рейтинг, текст, фото, привязка к товару и пользователю)'),
        ('News', 'новости (заголовок, slug, содержимое, теги, статус публикации)'),
        ('ChatSession', 'сессии чата (привязка к пользователю/гостю, статус)'),
        ('ChatMessage', 'сообщения чата (тип отправителя, содержимое, вложения)'),
        ('PromoCode', 'промокоды (код, email получателя, срок действия, статус использования)'),
        ('ContactMessage', 'сообщения из обратной связи (имя, email, телефон, тема, сообщение)'),
        ('AuditLog', 'журнал аудита действий пользователей'),
    ]
    
    for name, desc in db_models:
        add_paragraph(doc, f'– {name} — {desc}')
    
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 14 – Схема базы данных веб-сервиса «Славянка»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Основные связи между моделями:')
    relations = [
        'User 1:N Order — один пользователь может иметь много заказов;',
        'User 1:N Review — один пользователь может оставить много отзывов;',
        'Category 1:N Product — одна категория содержит много товаров;',
        'Product N:N Order (через OrderItem) — товар может входить в разные заказы;',
        'Product 1:N Review — один товар может иметь много отзывов;',
        'ChatSession 1:N ChatMessage — одна сессия содержит много сообщений.',
    ]
    for r in relations:
        add_paragraph(doc, f'– {r}')
    
    add_subsubheading(doc, '2.1.3 Алгоритмы обработки информации', size=14)
    
    add_paragraph(doc, 'В данном разделе описаны ключевые алгоритмы работы веб-сервиса.')
    
    add_paragraph(doc, 'Алгоритм работы корзины:', bold=True)
    add_paragraph(doc, 'При добавлении товара в корзину выполняется проверка его наличия в текущем списке. Если товар уже есть, увеличивается его количество, иначе создаётся новая позиция. При изменении количества пересчитывается общая стоимость заказа. При оформлении заказа проверяется авторизация пользователя, после чего создаётся запись в базе данных со статусом «new».')
    
    add_paragraph(doc, 'Алгоритм авторизации:', bold=True)
    add_paragraph(doc, 'Пользователь вводит email и пароль. Сервер проверяет существование пользователя и сравнивает хеш пароля с помощью bcrypt. При успешной проверке генерируется JWT-токен, который возвращается клиенту. Токен содержит ID пользователя и роль. При каждом запросе к защищённым маршрутам выполняется проверка JWT-токена.')
    
    add_paragraph(doc, 'Алгоритм работы умного адреса:', bold=True)
    add_paragraph(doc, 'При вводе адреса доставки выполняется поиск по списку ближайших населённых пунктов (nearbySettlements.ts). Пользователю предлагаются варианты автодополнения. После выбора населённого пункта формируется полный адрес доставки.')
    
    add_paragraph(doc, 'Алгоритм работы чата:', bold=True)
    add_paragraph(doc, 'При открытии чата создаётся сессия (для авторизованных пользователей привязывается к userId, для гостей — к guestId). Сообщения сохраняются в базе данных и передаются через WebSocket-соединение. Администратор видит все активные сессии и может отвечать на сообщения.')
    
    add_paragraph(doc, 'Алгоритм обработки заказа:', bold=True)
    steps = [
        'Пользователь добавляет товары в корзину;',
        'Переходит к оформлению заказа;',
        'Заполняет контактные данные (имя, телефон, адрес);',
        'Выбирает способ оплаты;',
        'Подтверждает заказ;',
        'Сервер создаёт заказ со статусом «new»;',
        'Заказ отображается в админ-панели;',
        'Администратор может изменить статус (processing, completed, cancelled).',
    ]
    for i, s in enumerate(steps):
        add_paragraph(doc, f'{i+1}. {s}')
    
    doc.add_page_break()
    
    # 2.2
    add_subheading(doc, '2.2 Описание работы сайта', size=15)
    
    add_paragraph(doc, 'Веб-сервис «Славянка» реализован с использованием архитектуры monorepo. Клиентская часть запускается на порту 3000 с помощью Vite, серверная — на порту 3001 с помощью Fastify. После запуска пользователь открывает браузер и переходит по адресу http://localhost:3000.')
    
    add_paragraph(doc, 'Карта веб-сервиса представлена на листе формата А4, прилагаемом к дипломному проекту в соответствии с приложением Б.')
    add_paragraph(doc, 'Блок-схема веб-сервиса представлена на листе формата А4, прилагаемом к дипломному проекту.')
    add_paragraph(doc, 'Схема базы данных веб-сервиса представлена на листе формата А4 с шифром ДП 09.02.07 ИС-422 11.26 БД, прилагаемая к дипломному проекту.')
    
    # 2.3
    add_subheading(doc, '2.3 Инструкция работы с сайтом', size=15)
    
    add_subsubheading(doc, '2.3.1 Общие сведения', size=14)
    
    add_paragraph(doc, 'Каждая страница веб-сервиса состоит из трёх базовых элементов:')
    
    add_paragraph(doc, 'Шапка содержит: логотип (ссылка на главную страницу); меню со ссылками на страницы «Каталог», «Акции», «Новости», «О магазине», «Контакты»; кнопку поиска; кнопку входа/личного кабинета; значок корзины с количеством товаров.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 15 – Шапка веб-сервиса', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Основная часть: меняется в зависимости от страницы, на которой находится пользователь.')
    
    add_paragraph(doc, 'Подвал содержит: логотип (ссылка на главную страницу); меню со ссылками на информационные страницы; контактную информацию (адрес, телефон, режим работы); ссылку на политику конфиденциальности.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 16 – Подвал веб-сервиса', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Веб-сервис содержит следующие страницы:')
    pages = [
        '«Главная» — информация о магазине, преимущества, популярные товары, отзывы, карта;',
        '«Каталог» — все товары с фильтрацией по категориям, поиском и сортировкой;',
        '«Акции» — товары со скидками и специальными предложениями;',
        '«Новости» — новости магазина и актуальные события;',
        '«О магазине» — подробная информация о магазине и его истории;',
        '«Контакты» — контактная информация, схема проезда, форма обратной связи;',
        '«Корзина» — выбранные товары для оформления заказа;',
        '«Личный кабинет» — данные пользователя, история заказов, избранное.',
    ]
    for p in pages:
        add_paragraph(doc, f'– {p}')
    
    add_paragraph(doc, 'Главная страница содержит информацию о магазине, его преимуществах, популярные товары, отзывы клиентов и интерактивную карту.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 17 – Главная страница веб-сервиса', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Страница «Каталог» позволяет просматривать все доступные товары, фильтровать их по категориям, искать по названию и сортировать. При нажатии на товар открывается модальное окно с подробной информацией (цена, описание, состав, вес, отзывы).')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 18 – Страница «Каталог»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Страница «Корзина» отображает добавленные товары с возможностью изменения количества, удаления товаров и оформления заказа.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 19 – Страница «Корзина»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'На странице личного кабинета отображается информация об аккаунте, история заказов, избранные товары. В верхней части страницы находится кнопка для редактирования информации аккаунта.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 20 – Страница «Личный кабинет»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Страница «Новости» содержит список новостей и акций магазина. При нажатии на новость открывается страница с полным текстом.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 21 – Страница «Новости»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'В административной панели (доступна по адресу /admin) администратор может:')
    admin_items = [
        'управлять товарами (добавление, редактирование, удаление);',
        'управлять категориями товаров;',
        'просматривать и изменять статусы заказов;',
        'управлять пользователями;',
        'управлять новостями;',
        'просматривать и отвечать на сообщения в чате;',
        'просматривать и отвечать на сообщения из формы обратной связи.',
    ]
    for item in admin_items:
        add_paragraph(doc, f'– {item}')
    
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 22 – Административная панель', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 23 – Чат с администратором', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 24 – Страница «Контакты»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_subsubheading(doc, '2.3.2 Вызов и загрузка', size=14)
    
    add_paragraph(doc, 'Веб-сервис на время тестирования находится на локальном сервере разработчика. Для запуска необходимо выполнить команды:')
    add_paragraph(doc, 'Бэкенд (серверная часть): cd apps/backend && npm run dev (порт 3001)')
    add_paragraph(doc, 'Фронтенд (клиентская часть): cd apps/frontend && npm run dev (порт 3000)')
    add_paragraph(doc, 'После запуска компонентов системы веб-сервис становится доступным по адресу: http://localhost:3000.')
    
    add_subsubheading(doc, '2.3.3 Входные данные', size=14)
    
    add_paragraph(doc, 'В процессе работы с веб-сервисом пользователь вводит данные, необходимые для функционирования системы. К входным данным относятся:')
    input_data = [
        'регистрационные данные (email, пароль, имя);',
        'данные для авторизации (email, пароль);',
        'поисковые запросы и фильтры;',
        'данные при оформлении заказа (имя, телефон, адрес);',
        'текст отзывов;',
        'сообщения в чате;',
        'текст сообщений из формы обратной связи.',
    ]
    for d in input_data:
        add_paragraph(doc, f'– {d}')
    
    add_subsubheading(doc, '2.3.4 Выходные данные', size=14)
    
    add_paragraph(doc, 'Выходными данными веб-сервиса является информация, формируемая системой в результате обработки запросов:')
    output_data = [
        'информация о магазине и товарах;',
        'результаты поиска и фильтрации;',
        'информация о заказах в личном кабинете;',
        'уведомления и информационные сообщения;',
        'статистические данные в административной панели;',
        'история чата.',
    ]
    for d in output_data:
        add_paragraph(doc, f'– {d}')
    
    add_subsubheading(doc, '2.3.5 Сообщения программы', size=14)
    
    add_paragraph(doc, 'В процессе работы веб-сервиса пользователю могут отображаться информационные сообщения и сообщения об ошибках:')
    
    add_table_caption(doc, 'Таблица 8 – Виды диалоговых окон')
    
    table8 = doc.add_table(rows=7, cols=2)
    table8.alignment = WD_TABLE_ALIGNMENT.CENTER
    table8.style = 'Table Grid'
    set_cell_text(table8.cell(0, 0), 'Вид сообщений', bold=True)
    set_cell_text(table8.cell(0, 1), 'Действия пользователя', bold=True)
    
    messages = [
        ('Ошибка авторизации', 'Проверить правильность email и пароля'),
        ('Незаполненные обязательные поля', 'Заполнить все обязательные поля формы'),
        ('Некорректное заполнение формы', 'Проверить введённые данные'),
        ('Успешное выполнение действия', 'Продолжить работу с сервисом'),
        ('Ошибка на стороне сервера', 'Повторить попытку позже'),
        ('Товар добавлен в корзину', 'Перейти в корзину или продолжить покупки'),
    ]
    for i, (msg, action) in enumerate(messages):
        set_cell_text(table8.cell(i+1, 0), msg, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table8.cell(i+1, 1), action, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_empty_line(doc)
    
    # 2.4
    add_subheading(doc, '2.4 Описание процесса отладки программы', size=15)
    
    add_subsubheading(doc, '2.4.1 Методы отладки', size=14)
    
    add_paragraph(doc, 'Отладка представляет собой процесс поиска, анализа и устранения ошибок, возникающих при разработке программного продукта. В процессе разработки веб-сервиса «Славянка» отладка проводилась на всех этапах создания.')
    
    add_paragraph(doc, 'Использовались следующие методы отладки:')
    debug_methods = [
        'Ручное тестирование — проверка функциональности путём взаимодействия с интерфейсом;',
        'Инструменты разработчика браузера (DevTools) — анализ сетевых запросов, консольных ошибок, состояния приложения;',
        'Логирование сервера Fastify — отслеживание ошибок серверной части;',
        'Prisma Studio — визуальная проверка состояния базы данных;',
        'TypeScript компилятор — выявление ошибок типизации на этапе разработки.',
    ]
    for m in debug_methods:
        add_paragraph(doc, f'– {m}')
    
    add_paragraph(doc, 'Основное внимание уделялось проверке корректности взаимодействия между клиентской и серверной частями, обработке данных форм, работе корзины и оформлению заказов.')
    
    add_subsubheading(doc, '2.4.2 Тестирование и отладка', size=14)
    
    add_paragraph(doc, 'В процессе разработки использовались различные средства тестирования:')
    
    add_table_caption(doc, 'Таблица 9 – Средства тестирования и отладки')
    
    table9 = doc.add_table(rows=7, cols=2)
    table9.alignment = WD_TABLE_ALIGNMENT.CENTER
    table9.style = 'Table Grid'
    set_cell_text(table9.cell(0, 0), 'Средство', bold=True)
    set_cell_text(table9.cell(0, 1), 'Назначение', bold=True)
    
    tools = [
        ('DevTools браузера', 'Проверка интерфейса, анализ сетевых запросов, консоль'),
        ('Консоль браузера', 'Просмотр ошибок JavaScript'),
        ('VS Code', 'Разработка, отладка TypeScript, интеграция с Git'),
        ('Логи Fastify', 'Отслеживание серверных ошибок'),
        ('Prisma Studio', 'Просмотр и редактирование данных в БД'),
        ('TypeScript', 'Статический анализ кода, выявление ошибок типов'),
    ]
    for i, (tool, desc) in enumerate(tools):
        set_cell_text(table9.cell(i+1, 0), tool, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table9.cell(i+1, 1), desc, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_subsubheading(doc, '2.4.3 Контрольный пример', size=14)
    
    add_paragraph(doc, 'В ходе выполнения дипломного проектирования была разработана поставленная задача — веб-сервис для интернет-магазина продуктов «Славянка». Рассмотрим процесс работы пользователя с веб-сервисом.')
    
    add_paragraph(doc, '1. Пользователь открывает главную страницу. На главной странице отображаются приветственный баннер, преимущества магазина, популярные товары, отзывы клиентов и интерактивная карта с адресом.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 25 – Главная страница', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '2. Пользователь переходит на страницу каталога через меню навигации. На странице отображаются все товары, разделённые по категориям. Пользователь может выбрать категорию, воспользоваться поиском или отсортировать товары.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 26 – Страница «Каталог»', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '3. Пользователь нажимает на товар. Открывается модальное окно с подробной информацией: название, цена, описание, состав, вес, срок годности, пищевая ценность, отзывы.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 27 – Модальное окно товара', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '4. Пользователь нажимает кнопку «В корзину». Товар добавляется, счётчик корзины обновляется.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 28 – Добавление товара в корзину', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '5. Пользователь переходит в корзину. Видит список товаров с ценами, может изменить количество или удалить товар. Нажимает «Оформить заказ».')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 29 – Корзина с товарами', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '6. Если пользователь не авторизован, система предлагает войти или зарегистрироваться. После авторизации пользователь заполняет данные для заказа.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 30 – Окно авторизации', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '7. Пользователь подтверждает заказ. Заказ отправляется администратору, отображается в личном кабинете.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 31 – Подтверждение заказа', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '8. Администратор в панели управления видит новый заказ, может изменить его статус.')
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 32 – Список заказов в админке', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 33 – Детали заказа в админке', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'Рисунок 34 – Чат с клиентом', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    doc.add_page_break()
    
    # ============================================================
    # 3 ЭКОНОМИЧЕСКАЯ ЧАСТЬ
    # ============================================================
    add_heading_centered(doc, '3 ЭКОНОМИЧЕСКАЯ ЧАСТЬ', size=16)
    
    add_paragraph(doc, 'Для оценки экономической эффективности проекта следует рассчитать полную себестоимость проекта и срок окупаемости с учётом маркетинговых исследований и рентабельности отрасли. Применяются методики по расчёту себестоимости по трудозатратам и показателям экономической эффективности.')
    
    add_paragraph(doc, 'Для подведения итогов в экономической части применяются различные формулы.')
    
    add_paragraph(doc, '1. Расчёт общих затрат', bold=True)
    add_empty_line(doc)
    add_paragraph(doc, 'Cобщ = Cразраб + Cоборуд + Cлицензии\t\t\t(1)', bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'где Cразраб — затраты на оплату труда разработчиков, Cоборуд — стоимость оборудования и инфраструктуры, Cлицензии — покупка лицензий на ПО и инструменты разработки.')
    
    add_paragraph(doc, '2. Расчёт годовой экономии', bold=True)
    add_empty_line(doc)
    add_paragraph(doc, 'Eгод = Eдо − Eпосле\t\t\t(2)', bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    add_paragraph(doc, 'где Eдо — расходы организации до внедрения проекта, Eпосле — ожидаемые расходы после внедрения.')
    
    add_paragraph(doc, '3. Расчётный срок окупаемости проекта', bold=True)
    add_empty_line(doc)
    add_paragraph(doc, 'Tокуп = Cобщ / Eгод\t\t\t(3)', bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, '4. Годовая экономия (коэффициент)', bold=True)
    add_empty_line(doc)
    add_paragraph(doc, 'Kэко = (Eгод / Cобщ) × 100%\t\t\t(4)', bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)
    
    add_paragraph(doc, 'После рассмотрения основных формул, применяемых для расчётов, необходимо применить значения, используемые для вычислений.')
    
    add_paragraph(doc, 'Первым этапом будет расчёт общих затрат на проект, для этого необходимо указать основные показатели:')
    
    add_paragraph(doc, 'Cразраб — это зарплата разработчика, рассчитывается по формуле:')
    add_paragraph(doc, 'Cразраб = ЗПе × количество месяцев работы\t\t\t(5)')
    add_paragraph(doc, 'где ЗПе — ежемесячная зарплата разработчика.')
    add_paragraph(doc, 'Cразраб = 45000 × 3 = 135000 руб.')
    
    add_paragraph(doc, 'Cоборуд — стоимость оборудования и инфраструктуры, включая аренду серверов.')
    add_paragraph(doc, 'Аренда сервера для программы — 15000 руб.')
    add_paragraph(doc, 'Итого Cоборуд = 15000 руб.')
    
    add_paragraph(doc, 'Cлицензии — покупка лицензий на ПО и инструменты разработки.')
    add_paragraph(doc, 'Лицензия Visual Studio 2026 — 12000 руб.')
    add_paragraph(doc, 'Лицензия SSMS — 8000 руб.')
    add_paragraph(doc, 'Итого Cлицензии = 20000 руб.')
    
    add_paragraph(doc, 'Получается, что расчёт общих затрат равен:')
    add_paragraph(doc, 'Cобщ = 135000 + 15000 + 20000 = 170000 руб.', bold=True)
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Расчёт годовой экономии производится в следующем порядке.', bold=True)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Эгод = Ссотруд + Соборуд + Сматериал\t\t\t(6)')
    add_paragraph(doc, 'где Эгод — экономия в результате внедрения программы, Ссотруд — зарплата сотрудника после автоматизации, Соборуд — стоимость оборудования после автоматизации, Сматериал — стоимость материалов после автоматизации.')
    
    add_paragraph(doc, 'Допустим, раньше компания ежегодно тратила на обработку документов вручную около 1 млн рублей, а внедрение автоматизированной системы снизило этот показатель до 600 тыс. рублей. Таким образом, годовая экономия составит:')
    add_paragraph(doc, 'Eгод = 1 000 000 − 600 000 = 400 000 руб.', bold=True)
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Для срока окупаемости подставляются готовые значения:')
    add_paragraph(doc, 'Tокуп = 170000 / 400000 = 2,35 года', bold=True)
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Далее производится расчёт годового экономического коэффициента:')
    add_paragraph(doc, 'Кэко = (400000 / 170000) × 100% = 23,53%', bold=True)
    
    add_empty_line(doc)
    
    add_paragraph(doc, 'Таким образом, внедрение проекта обеспечит ежегодную экономию в размере около 23,53% от первоначальных инвестиций.')
    
    add_empty_line(doc)
    add_empty_line(doc)
    
    add_table_caption(doc, 'Таблица 10 – Экономические показатели проекта')
    
    table10 = doc.add_table(rows=5, cols=2)
    table10.alignment = WD_TABLE_ALIGNMENT.CENTER
    table10.style = 'Table Grid'
    set_cell_text(table10.cell(0, 0), 'Наименование', bold=True)
    set_cell_text(table10.cell(0, 1), 'Величина, тыс. руб.', bold=True)
    
    econ_data = [
        ('1. Годовая экономия, руб.', '400'),
        ('2. Дополнительные капитальные вложения', '170'),
        ('3. Расчётный коэффициент экономической эффективности', '23,53'),
        ('4. Срок окупаемости, год', '2,35'),
    ]
    for i, (name, val) in enumerate(econ_data):
        set_cell_text(table10.cell(i+1, 0), name, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table10.cell(i+1, 1), val, size=11)
    
    add_empty_line(doc)
    add_empty_line(doc)
    
    add_paragraph(doc, 'Вывод: рассчитанный срок окупаемости проекта составляет 2,35 года, что входит в установленную норму (до 5 лет). Общие затраты на проект составили 170 000 рублей. Проект является экономически эффективным и целесообразным для внедрения. Годовая экономия после внедрения системы автоматизации составит 400 000 рублей, а коэффициент экономической эффективности — 23,53%.')
    
    doc.add_page_break()
    
    # ============================================================
    # ЗАКЛЮЧЕНИЕ
    # ============================================================
    add_heading_centered(doc, 'ЗАКЛЮЧЕНИЕ', size=16)
    
    add_paragraph(doc, 'В ходе дипломного проектирования был разработан веб-сервис «Славянка» для онлайн-продажи продуктов питания, а также для предоставления информации о магазине, его деятельности и ассортименте.')
    
    add_paragraph(doc, 'В ходе выполнения дипломного проекта были изучены современные технологии веб-разработки, проанализированы особенности проектирования клиентской и серверной частей приложения, рассмотрены принципы организации систем онлайн-торговли.')
    
    add_paragraph(doc, 'Проектное решение реализовано с помощью языка разметки HTML5 в связке с CSS3, TypeScript и JavaScript. Для клиентской части использован фреймворк Vue 3 с библиотеками Pinia и Vue Router. Для серверной части — фреймворк Fastify. В качестве ORM используется Prisma, база данных — SQLite.')
    
    add_paragraph(doc, 'Веб-сервис предоставляет пользователям следующие возможности:')
    final_features = [
        'просмотр каталога товаров с поиском, фильтрацией и сортировкой;',
        'регистрация и авторизация пользователей;',
        'добавление товаров в корзину и оформление заказа;',
        'просмотр истории заказов в личном кабинете;',
        'добавление товаров в избранное;',
        'просмотр новостей и акций магазина;',
        'написание отзывов на товары;',
        'чат с администратором магазина;',
        'использование промокодов;',
        'управление товарами, категориями, заказами, пользователями, новостями и чатами через административную панель.',
    ]
    for f in final_features:
        add_paragraph(doc, f'– {f}')
    
    add_paragraph(doc, 'Данные подгружаются на веб-сервис динамически без перезагрузки страницы за счёт технологии SPA. Интерфейс веб-сервиса интуитивно понятный и простой в использовании. Переходы между страницами и другие действия пользователя сопровождаются плавными анимациями.')
    
    add_paragraph(doc, 'В ходе тестирования была выполнена проверка работоспособности основных функциональных модулей системы. Проведённое тестирование подтвердило корректность работы разработанного программного продукта.')
    
    add_paragraph(doc, 'В экономической части была выполнена оценка эффективности разработки. Общие затраты на проект составили 170 000 рублей. Срок окупаемости — 2,35 года, что входит в нормативный период (5 лет). Проект является экономически целесообразным.')
    
    add_paragraph(doc, 'Таким образом, все поставленные задачи дипломного проекта были успешно решены. Во-первых, выполнен анализ предметной области, определены требования к веб-сервису и обоснован выбор технологий разработки. Во-вторых, разработана архитектура приложения, реализованы клиентская и серверная части, описаны алгоритмы работы программного продукта. В-третьих, проведено тестирование системы, подготовлено руководство пользователя и выполнено экономическое обоснование разработки. Поставленная цель — разработка веб-сервиса для интернет-магазина продуктов «Славянка», обеспечивающего возможность онлайн-просмотра каталога товаров, оформления заказов и управления ими — полностью достигнута.')
    
    add_paragraph(doc, 'Разработанный веб-сервис может использоваться в качестве основы для дальнейшего развития и внедрения дополнительных функциональных возможностей системы. В дальнейшем функционал веб-сервиса может быть расширен: добавление онлайн-оплаты, интеграция с внешними платёжными системами, мобильное приложение, система лояльности, интеграция с CRM.')
    
    # ============================================================
    # СПИСОК ЛИТЕРАТУРЫ
    # ============================================================
    add_heading_centered(doc, 'Список литературы', size=16)
    
    literature = [
        'ГОСТ 2.105-95 Общие требования к текстовым документам. – М.: Издательство стандартов, 1994.',
        'ГОСТ 19.001-77 Единая система программной документации. Общие положения. – М.: Издательство стандартов, 1995.',
        'ГОСТ 19.002-80 Единая система программной документации. Схемы алгоритмов и программ. Правила выполнения. – М.: Издательство стандартов, 1995.',
        'ГОСТ 19.103-77 Единая система программной документации. Обозначение программ и программных документов. – М.: Издательство стандартов, 1994.',
        'ГОСТ 19.104-78 Единая система программной документации. Основные надписи. – М.: Издательство стандартов, 1995.',
        'ГОСТ 19.105-78 Единая система программной документации. Общие требования к программным документам. – М.: Издательство стандартов, 1995.',
        'ГОСТ 19.402-78 Единая система программной документации. Описание программы. – М.: Издательство стандартов, 1996.',
        'ГОСТ 19.404-79 Единая система программной документации. Пояснительная записка. Требования к содержанию и оформлению. – М.: Издательство стандартов, 1994.',
        'Эрик Х. Vue.js в действии / Х. Эрик, Л. Бенджамин. – Санкт-Петербург: Питер, 2021. – 304 с.',
        'Никсон Р. – Создаем динамические веб-сайты с помощью PHP, MySQL, JavaScript, CSS и HTML5. 6-е издание / Р. Никсон. – Санкт-Петербург: Питер, 2023. – 832 с.',
        'Дакетт Д. HTML и CSS. Разработка и дизайн веб-сайтов / Д. Дакетт. – М.: Эксмо, 2023. – 480 с.',
        'Кириченко А.В. TypeScript. Полное руководство / А.В. Кириченко. – Санкт-Петербург: Наука и Техника, 2024. – 320 с.',
    ]
    
    for lit in literature:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(lit)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    doc.save(os.path.join(OUTPUT_DIR, '5 ПЗ.docx'))
    print('[OK] 5 ПЗ.docx')

# ============================================================
# ГЛАВНАЯ ФУНКЦИЯ
# ============================================================
def main():
    print('=' * 60)
    print('Генерация дипломных документов для проекта "Славянка"')
    print('=' * 60)
    
    ensure_output_dir()
    copy_static_files()
    
    generate_titulnik()
    generate_zadanie()
    generate_tz()
    generate_annotaciya()
    generate_pz()
    generate_vedomost()
    generate_spravka()
    generate_prilozhenie_a()
    generate_prilozhenie_b()
    generate_prilozhenie_v()
    
    print('=' * 60)
    print('Генерация завершена!')
    print('All files saved in folder:', OUTPUT_DIR)
    print('=' * 60)

if __name__ == '__main__':
    main()